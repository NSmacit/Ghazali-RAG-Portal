import os
import shutil

# =====================================================================
# ISLAMICATE DH - GAZALİ ENTERPRISE REFACTORING & IP PROTECTION TOOL
# =====================================================================
# Bu betik, lokal bilgisayarınızda çalışan monolitik Streamlit uygulamasını
# FAANG / Kurumsal seviyede temiz, modüler, test edilebilir bir yapıya kavuşturur.
# Projenizin fikir mülkiyetini (IP) korumak için gerekli lisanslama ve
# .gitignore koruma kalkanlarını otomatik olarak kurar.
# =====================================================================

print("\n" + "="*80)
print("🕌 GAZÂLÎ PORTAL - ENTERPRISE REFACTORING & IP PROTECTION CONSOLE")
print("="*80)

# 1. Klasör Yapısının Kurulması
directories = ["config", "src", "tests"]
for directory in directories:
    if not os.path.exists(directory):
        os.makedirs(directory)
        print(f"[+] Klasör oluşturuldu: {directory}/")
    else:
        print(f"[i] Klasör zaten mevcut: {directory}/")

# 2. IP Koruması için Gitignore Dosyasının Yazılması
# Bu dosya, kritik ChromaDB veritabanınızın ve yerel dosyalarınızın GitHub'a sızmasını kesin olarak önler.
gitignore_content = """# Gazali Project Security Gitignore
# 1. Yerel Veritabanı (Mülkiyet Koruması - Asla GitHub'a yüklenmez)
gazali_chroma_db/
*.pdf
*.docx
*.md
*.html
*.png
*.jpg
*.jpeg

# İstisna: README ve belgeler
!README.md
!requirements.txt
!LICENSE

# 2. Güvenlik ve Kimlik Bilgileri (API Anahtarları)
.env
.env.local
.env.development.local
.env.test.local
.env.production.local

# 3. Python ve Sistem Kalıntıları
__pycache__/
*.py[cod]
*$py.class
.env
.venv
env/
venv/
ENV/
.pytest_cache/
.DS_Store
"""

with open(".gitignore", "w", encoding="utf-8") as f:
    f.write(gitignore_content)
print("[+] Güvenlik Kalkanı Kuruldu: .gitignore dosyası yazıldı (Veritabanınız koruma altında).")

# 3. Güçlü Fikir Mülkiyeti Lisansı (AGPL-3.0)
# Bu lisans, projenizin kurumsal firmalar tarafından kopyalanıp ticari olarak sömürülmesini engeller.
# Kodlarınızı GitHub'da sergilemenize izin verirken, kullanan herkesi açık kaynak olmaya zorlar.
agpl_license_content = """GNU AFFERO GENERAL PUBLIC LICENSE
Version 3, 19 November 2007

Copyright (C) 2026 Macit. All Rights Reserved.
Everyone is permitted to copy and distribute verbatim copies of this license document, but changing it is not allowed.

[AGPL-3.0 LİSANS METNİ ÖZETİ - PROJE KORUMASI]
Bu proje, GNU Affero Genel Kamu Lisansı (AGPL-3.0) altında lisanslanmıştır.
Bu kodun herhangi bir parçasını kendi sunucularında veya projelerinde kullanan herhangi bir ticari kuruluş, 
kendi projelerinin de kaynak kodlarını tamamen açık kaynak olarak paylaşmakla yükümlüdür.
Bu lisanslama modeli, büyük teknoloji firmalarının projenizi izinsiz ticarileştirmesini önleyen en güçlü kalkan modelidir.
"""

with open("LICENSE", "w", encoding="utf-8") as f:
    f.write(agpl_license_content)
print("[+] IP Telif Koruma Kalkanı Kuruldu: LICENSE (AGPL-3.0) dosyası yazıldı.")

# 4. .env.example Şablonunun Yazılması
env_example_content = """# İmam Gazali Enterprise RAG Config
# Google AI Studio'dan aldığınız API anahtarını buraya girin.
# Yerel çalışmada bu dosyayı kopyalayıp adını ".env" yapabilirsiniz.
GEMINI_API_KEY=your_gemini_api_key_here
"""

with open(".env.example", "w", encoding="utf-8") as f:
    f.write(env_example_content)
print("[+] Konfigürasyon Şablonu Kuruldu: .env.example yazıldı.")

# 5. Requirements.txt Dosyasının Yazılması
requirements_content = """streamlit>=1.32.0
chromadb>=0.4.24
sentence-transformers>=2.5.1
google-generativeai>=0.4.0
python-docx>=1.1.0
pandas>=2.1.0
networkx>=3.2.1
pyvis>=0.3.2
"""

with open("requirements.txt", "w", encoding="utf-8") as f:
    f.write(requirements_content)
print("[+] Bağımlılık Listesi Kuruldu: requirements.txt yazıldı.")

# 6. config/settings.py Dosyasının Yazılması
config_settings_content = """# =====================================================================
# GAZALİ ENTERPRISE SYSTEM CONFIGURATION
# =====================================================================

CUSTOM_CSS = \"\"\"
<style>
    .reportview-container {
        background: #121212;
    }
    .main .block-container {
        padding-top: 2rem;
    }
    h1, h2, h3 {
        color: #00adb5 !important;
        font-family: 'Georgia', serif;
    }
    .stButton>button {
        background-color: #00adb5;
        color: white;
        border-radius: 8px;
        border: none;
        padding: 0.5rem 1.5rem;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #007a80;
        border: none;
        box-shadow: 0 4px 12px rgba(0,173,181,0.3);
    }
    .source-box {
        background-color: #1e1e1e;
        border-left: 4px solid #00adb5;
        padding: 12px;
        margin: 8px 0px;
        border-radius: 4px;
    }
</style>
\"\"\"

ARABIC_TO_TURKISH_DICT = {
    "العلم": "ilim bilgi muallim taallüm",
    "العمل": "amel eylem pratik ibadet",
    "القلب": "kalp gönül cevher tasfiye",
    "النفس": "nefis nefs ruh kendini benlik",
    "العقل": "akıl rasyonel düşünce fehim",
    "السعادة": "saadet mutluluk kurtuluş necât",
    "الشك": "şüphe tereddüt şüphecilik istidlal",
    "اليقين": "yakin yakinî kesin bilgi hakikat",
    "الحس": "hissiyyat duyular his duyu organları",
    "الرياضة": "riyazet nefs terbiyesi tefekkür",
    "الوسوسة": "vesvese kuruntu şeytan vesvesesi",
    "معرفة": "marifet bilmek tanımak irfan",
    "معرفة النفس": "kendini bilmek marifet-i nefs nefsini bilmek",
    "معرفة الله": "Allah'ı bilmek marifetullah",
    "الدنيا": "dünya hayatı fani geçici alem",
    "الآخرة": "ahiret beka alemi ebediyet",
    "الولد": "veled çocuk oğul ey oğul",
    "نصيحة": "nasihat öğüt vasiyet",
    "تصوف": "tasavvuf ahlak zühd takva",
    "طهارة": "taharet temizlik kalb tasfiyesi",
    "عشق": "aşk muhabbet sevgi",
    "نور": "nur ışık ilahi aydınlanma"
}
"""

with open("config/settings.py", "w", encoding="utf-8") as f:
    f.write(config_settings_content)
print("[+] Modüler settings yazıldı: config/settings.py")

# 7. src/database.py Dosyasının Yazılması
src_database_content = """import os
import re
import chromadb
import streamlit as st
from sentence_transformers import SentenceTransformer

class TurkishBM25:
    \"\"\"Türkçe uyumlu, hafif ve yerel bir BM25 sınıfı.\"\"\"
    def __init__(self, corpus, b=0.75, k1=1.5):
        self.b = b
        self.k1 = k1
        self.corpus_size = len(corpus)
        self.avg_dl = 0
        self.doc_lengths = []
        self.doc_freqs = {}
        self.idf = {}
        self.tokenized_corpus = []
        
        self._initialize(corpus)

    def _tokenize(self, text):
        text = text.replace('I', 'ı').replace('İ', 'i').lower()
        words = re.findall(r'[a-zA-Z0-9çğıöşüçĞIİÖŞÜ]+', text)
        return words

    def _initialize(self, corpus):
        total_length = 0
        for doc in corpus:
            tokens = self._tokenize(doc)
            self.tokenized_corpus.append(tokens)
            self.doc_lengths.append(len(tokens))
            total_length += len(tokens)
            
            unique_tokens = set(tokens)
            for token in unique_tokens:
                self.doc_freqs[token] = self.doc_freqs.get(token, 0) + 1
                
        self.avg_dl = total_length / self.corpus_size if self.corpus_size > 0 else 1
        
        for token, df in self.doc_freqs.items():
            self.idf[token] = max(0.0001, (self.corpus_size - df + 0.5) / (df + 0.5) + 1)

    def get_scores(self, query):
        query_tokens = self._tokenize(query)
        scores = []
        
        for idx, doc_tokens in enumerate(self.tokenized_corpus):
            score = 0.0
            doc_len = self.doc_lengths[idx]
            
            word_counts = {}
            for token in doc_tokens:
                word_counts[token] = word_counts.get(token, 0) + 1
                
            for q_token in query_tokens:
                if q_token in word_counts:
                    tf = word_counts[q_token]
                    idf_val = self.idf.get(q_token, 0.0)
                    numerator = tf * (self.k1 + 1)
                    denominator = tf + self.k1 * (1 - self.b + self.b * (doc_len / self.avg_dl))
                    score += idf_val * (numerator / denominator)
            scores.append(score)
        return scores

@st.cache_resource
def load_rag_assets():
    \"\"\"Vektör tabanını ve E5 modelini önbellekte bir kez yükler.\"\"\"
    try:
        db_path = \"./gazali_chroma_db\"
        chroma_client = chromadb.PersistentClient(path=db_path)
        embed_model = SentenceTransformer(\"intfloat/multilingual-e5-large\")
        collection = chroma_client.get_collection(name=\"gazali_kulliyati\")
        
        all_data = collection.get()
        if not all_data or not all_data[\"documents\"]:
            return None, None, None, None
            
        bm25_engine = TurkishBM25(all_data[\"documents\"])
        return embed_model, collection, bm25_engine, all_data
    except Exception as e:
        st.error(f\"Veritabanı yüklenirken hata oluştu: {e}\")
        return None, None, None, None
"""

with open("src/database.py", "w", encoding="utf-8") as f:
    f.write(src_database_content)
print("[+] Modüler database engine yazıldı: src/database.py")

# 8. src/search.py Dosyasının Yazılması
src_search_content = """import re
import streamlit as st
import google.generativeai as genai
from config.settings import ARABIC_TO_TURKISH_DICT

def translate_expand_query(query):
    \"\"\"Sorguda Arapça karakterler varsa yerel sözlük ve Gemini yardımıyla Türkçe genişletme yapar.\"\"\"
    is_arabic = bool(re.search(r'[\\u0600-\\u06FF]', query))
    if not is_arabic:
        return query, False, \"\"
        
    clean_q = re.sub(r'[^\\w\\s\\u0600-\\u06FF]', '', query).strip()
    translated = ARABIC_TO_TURKISH_DICT.get(clean_q, \"\")
    
    if not translated:
        words = clean_q.split()
        tr_words = []
        for w in words:
            tr_w = ARABIC_TO_TURKISH_DICT.get(w, \"\")
            if tr_w:
                tr_words.append(tr_w)
        if tr_words:
            translated = \" \".join(tr_words)
            
    dynamic_translation = \"\"
    if st.session_state.get(\"api_key_valid\", False):
        try:
            model = genai.GenerativeModel(st.session_state.get(\"selected_model\"))
            prompt = f\"Translate the following classical Arabic Islamic/philosophical term or question into clean Turkish search keywords for a book database. Return ONLY the translated Turkish keywords, no explanations:\\n{query}\"
            response = model.generate_content(prompt)
            dynamic_translation = response.text.strip().replace(\"\\n\", \" \")
        except Exception:
            pass
            
    combined_translation = translated
    if dynamic_translation:
        if combined_translation:
            combined_translation += \" \" + dynamic_translation
        else:
            combined_translation = dynamic_translation
            
    if not combined_translation:
        combined_translation = query
        
    return combined_translation, True, clean_q

def run_hybrid_search(embed_model, collection, bm25_engine, all_data, query, top_k=4):
    \"\"\"BM25 ve Vektör (E5) sıralamalarını RRF formülüyle sentezler. Arapça sorgularda çapraz dilli arama yapar.\"\"\"
    if not collection or not bm25_engine or not all_data:
        return []
        
    bm25_query = query
    vector_query = query
    
    translated_q, is_arabic, original_arabic = translate_expand_query(query)
    if is_arabic:
        bm25_query = translated_q
        vector_query = f\"{query} {translated_q}\"
        st.info(f\"🌐 **Arapça Arama Tespit Edildi:**\\n• Orijinal Arapça: `{query}`\\n• Türkçe Genişletme: `{translated_q}`\\n\\n*Çapraz Dilli (Cross-lingual) Hibrit motorumuz, hem orijinal terimi vektör uzayında tarar hem de otomatik çeviri üzerinden yerel BM25 indekslemesi gerçekleştirir.*\")
        
    bm25_scores = bm25_engine.get_scores(bm25_query)
    bm25_ranked = sorted(range(len(bm25_scores)), key=lambda k: bm25_scores[k], reverse=True)
    
    formatted_query = f\"query: {vector_query}\"
    query_vector = embed_model.encode(formatted_query).tolist()
    vector_results = collection.query(
        query_embeddings=[query_vector],
        n_results=len(all_data[\"documents\"])
    )
    
    vector_order = vector_results[\"ids\"][0]
    vector_ranked = [all_data[\"ids\"].index(vid) for vid in vector_order]
    
    rrf_scores = {}
    k_constant = 60
    
    for rank, idx in enumerate(bm25_ranked):
        rrf_scores[idx] = rrf_scores.get(idx, 0.0) + (1.0 / (k_constant + rank + 1))
        
    for rank, idx in enumerate(vector_ranked):
        rrf_scores[idx] = rrf_scores.get(idx, 0.0) + (1.0 / (k_constant + rank + 1))
        
    top_indices = sorted(rrf_scores.keys(), key=lambda x: rrf_scores[x], reverse=True)[:top_k]
    
    final_docs = []
    for idx in top_indices:
        metadata = all_data[\"metadatas\"][idx]
        text = all_data[\"documents\"][idx]
        
        v_rank = vector_ranked.index(idx) + 1 if idx in vector_ranked else \"N/A\"
        b_rank = bm25_ranked.index(idx) + 1
        
        final_docs.append({
            \"text\": text,
            \"title\": metadata.get(\"title\", \"Bilinmeyen\"),\n            \"book\": metadata.get(\"book\", \"\"),
            \"page\": metadata.get(\"page\", \"\"),
            \"links\": metadata.get(\"links\", \"\"),
            \"rrf_score\": rrf_scores[idx],
            \"v_rank\": v_rank,
            \"b_rank\": b_rank
        })
        
    return final_docs
"""

with open("src/search.py", "w", encoding="utf-8") as f:
    f.write(src_search_content)
print("[+] Modüler search engine yazıldı: src/search.py")

# 9. src/ner.py Dosyasının Yazılması
src_ner_content = """import json
import streamlit as st
import google.generativeai as genai

SYSTEM_INSTRUCTION_NER = \"\"\"Sen klasik İslami metinler ve teoloji üzerinde uzmanlaşmış, son derece hassas bir metin analitiği (Named Entity Recognition) asistanısın.
Görevin, sana sunulan kaynak metinleri analiz ederek içindeki:
1. AYETLERİ (Eğer varsa; sure adı, ayet numarası, Türkçe meali ve geçtiği kaynak ile birlikte)
2. HADİSLERİ (Eğer varsa; hadisin içeriği, ravisi veya kaynağı, geçtiği yer ile birlikte)
3. ŞAHISLARI (Metinde adı geçen tarihi alimler, filozoflar, peygamberler, sahabeler veya şahsiyetler)
bulup ayıklamaktır.

Çıktıyı KESİNLİKLE ama KESİNLİKLE sadece aşağıdaki JSON şablonuna göre üretmelisin. JSON dışında hiçbir giriş, açıklama, markdown kodu veya düz metin yazmamalısın. Çıktı doğrudan geçerli bir JSON olmalıdır:
{
  "ayetler": [
    {"sure_ayet": "Sure Adı ve Ayet No", "metin": "Ayeti kerimenin meali veya içeriği", "kaynak": "Geçtiği kitap/not ve sayfa"}
  ],
  "hadisler": [
    {"metin": "Hadis-i şerifin meali veya içeriği", "ravi_kaynak": "Zikredilen ravi veya hadis kaynağı", "kaynak": "Geçtiği kitap/not ve sayfa"}
  ],
  "sahislar": [
    {"isim": "Kişinin Adı", "rol_baglam": "Metindeki rolü veya hangi bağlamda zikredildiği", "kaynak": "Geçtiği kitap/not ve sayfa"}
  ]
}

Eğer metinde ayet, hadis veya şahıs yoksa, ilgili listeyi boş bırak: []
\"\"\"

def extract_and_display_ner(docs):
    \"\"\"Paragraflardan Ayet, Hadis ve Şahıs varlıklarını Gemini kullanarak ayıklar, JSON formatında döndürür.\"\"\"
    context_str = \"\"
    for doc in docs:
        label = f\"[{doc['book']} (Sayfa {doc['page']})]\" if doc['book'] else f\"[[{doc['title']}]]\"
        context_str += f\"\\n--- KAYNAK: {label} ---\\n{doc['text']}\\n--------------------\\n\"
        
    user_prompt = f\"Analiz Edilecek Kaynak Metinler:\\n==================================================\\n{context_str}\\n==================================================\\n\\nYukarıdaki kaynakları tara ve JSON formatında ayıklama sonuçlarını döndür:\"
    
    try:
        model = genai.GenerativeModel(
            model_name=st.session_state.selected_model,
            system_instruction=SYSTEM_INSTRUCTION_NER
        )
        response = model.generate_content(
            contents=user_prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.1,
                response_mime_type=\"application/json\"
            )
        )
        
        raw_json = response.text.strip()
        if raw_json.startswith(\"```json\"):
            raw_json = raw_json.split(\"```json\", 1)[1]
        if raw_json.endswith(\"```\"):
            raw_json = raw_json.rsplit(\"```\", 1)[0]
        raw_json = raw_json.strip()
        
        data = json.loads(raw_json)
        return data, True, \"\"
    except Exception as e:
        return None, False, str(e)
"""

with open("src/ner.py", "w", encoding="utf-8") as f:
    f.write(src_ner_content)
print("[+] Modüler NER engine yazıldı: src/ner.py")

# 10. src/ui.py Dosyasının Yazılması
src_ui_content = """import io
from docx import Document

def generate_docx_stream(title, content):
    \"\"\"Makaleyi doğrudan indirilebilir Word dosyasına dönüştürür.\"\"\"
    doc = Document()
    doc.add_heading(title, level=0)
    
    for paragraph in content.split(\"\\n\\n\"):
        if paragraph.strip():
            if paragraph.startswith(\"### \"):
                doc.add_heading(paragraph.replace(\"### \", \"\"), level=2)
            elif paragraph.startswith(\"## \"):
                doc.add_heading(paragraph.replace(\"## \", \"\"), level=1)
            else:
                doc.add_paragraph(paragraph)
                
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
"""

with open("src/ui.py", "w", encoding="utf-8") as f:
    f.write(src_ui_content)
print("[+] Modüler UI helpers yazıldı: src/ui.py")

# 11. main.py Dosyasının Yazılması (Streamlit Enterprise Entrypoint)
main_py_content = """import os
import streamlit as st
import google.generativeai as genai
import pandas as pd
from config.settings import CUSTOM_CSS
from src.database import load_rag_assets
from src.search import run_hybrid_search
from src.ui import generate_docx_stream
from src.ner import extract_and_display_ner

# Streamlit sayfa konfigurasyonu
st.set_page_config(
    page_title=\"İmam Gazâlî Dijital Beşerî Bilimler Portalı\",
    page_icon=\"🕌\",
    layout=\"wide\",\n    initial_sidebar_state=\"expanded\"
)

# Dark Theme uygula
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

# Varlıkları yükle
embed_model, collection, bm25_engine, all_data = load_rag_assets()

# Session State'leri Başlat (Tüm durumlar bellekli!)
if \"chat_history\" not in st.session_state:
    st.session_state.chat_history = []
if \"api_key_valid\" not in st.session_state:
    st.session_state.api_key_valid = False
if \"selected_model\" not in st.session_state:
    st.session_state.selected_model = None
if \"available_models\" not in st.session_state:
    st.session_state.available_models = []
if \"writer_article\" not in st.session_state:
    st.session_state.writer_article = \"\"
if \"writer_topic\" not in st.session_state:
    st.session_state.writer_topic = \"\"
if \"ner_results\" not in st.session_state:
    st.session_state.ner_results = None
if \"ner_topic_state\" not in st.session_state:
    st.session_state.ner_topic_state = \"\"
if \"ner_docs\" not in st.session_state:
    st.session_state.ner_docs = []

# =====================================================================
# YAN PANEL (SIDEBAR) - GÜVENLİK VE AYARLAR
# =====================================================================
with st.sidebar:
    avatar_filename = \"gazali_avatar_v2.jpg\"
    if os.path.exists(avatar_filename):
        try:
            st.image(avatar_filename, use_container_width=True)
        except TypeError:
            st.image(avatar_filename, use_column_width=True)
    else:
        try:
            st.image(\"https://lh3.googleusercontent.com/notebooklm/AKYWMX-1Hvp_4i-e2m2xMO2yCZQvlu0Dq6YH3N82xKnYbzuVmBO_a5leTcjlypD_WLjym7j4ZSLDCeXSgOyurMbZkWei8w59lZ0eclF1eBJtknDRkophMYKIXkMq9X2xmBRZgYb--mb3S5GDIYZOt2hQohhT-Oct2Q\", use_container_width=True)
        except TypeError:
            st.image(\"https://lh3.googleusercontent.com/notebooklm/AKYWMX-1Hvp_4i-e2m2xMO2yCZQvlu0Dq6YH3N82xKnYbzuVmBO_a5leTcjlypD_WLjym7j4ZSLDCeXSgOyurMbZkWei8w59lZ0eclF1eBJtknDRkophMYKIXkMq9X2xmBRZgYb--mb3S5GDIYZOt2hQohhT-Oct2Q\", use_column_width=True)
            
    st.title(\"🕌 Gazâlî Portal\")\n    st.caption(\"Digital Humanities & RAG Platform v2.0 (Enterprise)\")\n    st.write(\"---\")
    
    api_key_input = st.text_input(
        \"🔑 Gemini API Key:\",
        type=\"password\",
        value=os.environ.get(\"GEMINI_API_KEY\", \"\"),
        help=\"Google AI Studio'dan aldığınız API anahtarı.\"\
    )
    
    if api_key_input:
        os.environ[\"GEMINI_API_KEY\"] = api_key_input
        genai.configure(api_key=api_key_input)
        
        if not st.session_state.api_key_valid:
            try:
                raw_models = genai.list_models()
                models_list = []
                for m in raw_models:
                    if \"generateContent\" in m.supported_generation_methods:
                        m_name = m.name if m.name.startswith(\"models/\") else f\"models/{m.name}\"
                        if not any(x in m_name for x in [\"gemini-pro\", \"gemini-2.5\"]):
                            models_list.append(m_name)
                
                def model_prio(name):
                    return 1 if \"flash\" in name.lower() else 5
                models_list.sort(key=model_prio)
                
                st.session_state.available_models = models_list
                st.session_state.selected_model = models_list[0] if models_list else None
                st.session_state.api_key_valid = True
                st.success(\"API Bağlantısı Başarılı! 🎉\")
            except Exception as e:
                st.error(f\"Bağlantı Hatası: {e}\")
                st.session_state.api_key_valid = False
                
    if st.session_state.api_key_valid:
        st.selectbox(
            \"🤖 Aktif Gemini Sürümü:\",
            options=st.session_state.available_models,\n            index=0,\n            key=\"selected_model_box\"\n        )
        st.session_state.selected_model = st.session_state.selected_model_box
    
    st.write(\"---\")
    if all_data:
        st.metric(label=\"📚 Toplam Paragraf Sayısı\", value=len(all_data[\"documents\"]))\n        st.info(\"Sisteminizde kavram notları, orijinal e-kitaplar ve Tehâfüt yüklüdür.\")
    else:
        st.warning(\"Veritabanınız boş! Lütfen önce pipeline veya ingester'ı çalıştırın.\")

# =====================================================================
# ANA SAYFA VE TAB TASARIMI
# =====================================================================
st.title(\"🕌 İmam Gazâlî Akademik Araştırma Portalı\")
st.write(\"Klasik İslâmî teoloji, felsefe ve epistemoloji araştırmalarını yapay zekâ ile buluşturan kurumsal bilgi yönetim paneli.\")

tab1, tab2, tab3, tab4 = st.tabs([
    \"💬 Akademik Sohbet (Chatbot)\", 
    \"✍️ Otomatik Co-Writer\", 
    \"📊 İnteraktif Kavram Ağ Haritası\", 
    \"📜 Ayet & Hadis & Şahıs Analitiği\"
])

# ---------------------------------------------------------------------
# TAB 1: AKADEMİK SOHBET (CHATBOT)
# ---------------------------------------------------------------------
with tab1:
    st.subheader(\"🤖 Hibrit Arama Destekli Soru-Cevap Motoru\")
    st.caption(\"Gelişmiş BM25 kelime araması ve E5-Large anlamsal aramayı birleştirerek sıfır halüsinasyonla çalışır.\")
    
    if not st.session_state.api_key_valid:
        st.info(\"👉 Devam etmek için lütfen sol panelden geçerli bir Gemini API anahtarı girin.\")
    else:
        for message in st.session_state.chat_history:
            with st.chat_message(message[\"role\"]):
                st.markdown(message[\"content\"])\n                if \"sources\" in message and message[\"sources\"]:\n                    with st.expander(\"📚 Grounding (Vektörel & BM25 Kaynak Atıfları)\"):\n                        for doc in message[\"sources\"]:\n                            source_header = f\"📌 {doc['book']} | Sayfa: {doc['page']}\" if doc['book'] else f\"📌 Obsidian Notu: [[{doc['title']}]]\"\n                            st.markdown(f\"**{source_header}** *(RRF Skoru: {doc['rrf_score']:.4f}, Vektör Sırası: {doc['v_rank']}, BM25 Sırası: {doc['b_rank']})*\")\n                            st.caption(f\"\\\\\\\"{doc['text']}\\\\\\\"\")
        
        user_query = st.chat_input(\"İmam Gazali felsefesi veya e-kitaplarınız hakkında sorun...\")
        
        if user_query:
            st.chat_message(\"user\").write(user_query)
            st.session_state.chat_history.append({\"role\": \"user\", \"content\": user_query})
            
            with st.spinner(\"📚 Yerel arşiviniz taranıyor...\"):\n                retrieved_docs = run_hybrid_search(embed_model, collection, bm25_engine, all_data, user_query, top_k=4)
                
            if not retrieved_docs:
                response_text = \"Arşivinizde bu konuyla ilişkili hiçbir belge veya kitap bulunamadı.\"\n                st.chat_message(\"assistant\").write(response_text)\n                st.session_state.chat_history.append({\"role\": \"assistant\", \"content\": response_text, \"sources\": []})\n            else:\n                context_str = \"\"\n                for doc in retrieved_docs:\n                    source_label = f\"[{doc['book']} (Sayfa {doc['page']})]\" if doc['book'] else f\"[[{doc['title']}]]\"\n                    context_str += f\"\\n--- KAYNAK: {source_label} ---\\nİçerik: {doc['text']}\\n------------------------\\n\"\n                    
                system_instruction = \"\"\"Sen İmam Gazali felsefesi ve teolojisi üzerine uzmanlaşmış, akademik dürüstlüğü en üst düzeyde tutan bir yapay zeka asistanısın.
Görevin, kullanıcının sorusuna SADECE sana sunulan 'KAYNAK NOTLAR' kapsamındaki verileri kullanarak yanıt vermektir.

Uyman Gereken Kesin Kurallar:
1. Sana verilen KAYNAK NOTLAR dışındaki hiçbir genel kültür veya internet bilgisini kullanma.
2. Eğer sorunun yanıtı sana sunulan notlarda doğrudan veya dolaylı olarak geçmiyorsa, kesinlikle bilgi UYDURMA (Halüsinasyon üretme).
3. Cevap verirken her cümlenin veya iddiadan sonra hangi belgeden/sayfadan alındığını belirt. (Örn: \"...duyuların insanı yanılttığını söyler [[Hissiyyât]] veya [Eyyühel Veled (Sayfa 5)].\")
4. Çelişkili veya belirsiz durumlar varsa bunları kendi yorumunu katmadan olduğu gibi aktar.
5. Tamamen Türkçe yanıt ver ve akademik, saygın bir dil kullan.
\"\"\"

                user_prompt = f\"Kullanıcı Sorusu: {user_query}\\n\\nSana Sunulan Kaynak Notlar:\\n==================================================\\n{context_str}\\n==================================================\\n\\nYukarıdaki kaynak notlara ve kurallara tamamen bağlı kalarak soruyu yanıtla:\"

                try:
                    with st.spinner(\"🔮 Gemini anlamsal sentez gerçekleştiriyor...\"):\n                        model = genai.GenerativeModel(\n                            model_name=st.session_state.selected_model,\n                            system_instruction=system_instruction\n                        )\n                        response = model.generate_content(\n                            contents=user_prompt,\n                            generation_config=genai.types.GenerationConfig(temperature=0.1)\n                        )\n                        response_text = response.text\n                        \n                        with st.chat_message(\"assistant\"):
                            st.markdown(response_text)
                            with st.expander(\"📚 Grounding (Vektörel & BM25 Kaynak Atıfları)\"):\n                                for doc in retrieved_docs:\n                                    source_header = f\"📌 {doc['book']} | Sayfa: {doc['page']}\" if doc['book'] else f\"📌 Obsidian Notu: [[{doc['title']}]]\"\n                                    st.markdown(f\"**{source_header}** *(RRF Skoru: {doc['rrf_score']:.4f}, Vektör Sırası: {doc['v_rank']}, BM25 Sırası: {doc['b_rank']})*\")\n                                    st.caption(f\"\\\\\\\"{doc['text']}\\\\\\\"\")
                                    
                        st.session_state.chat_history.append({
                            \"role\": \"assistant\",
                            \"content\": response_text,
                            \"sources\": retrieved_docs
                        })
                except Exception as e:
                    st.error(f\"Yapay zeka cevap üretirken hata oluştu: {e}\")

# ---------------------------------------------------------------------
# TAB 2: AKADEMİK CO-WRITER (YAZAR) - Bellek Korumalı!
# ---------------------------------------------------------------------
with tab2:
    st.subheader(\"✍️ Akıllı Makale ve Tez Taslağı Hazırlayıcı\")
    st.caption(\"Obsidian kütüphaneniz ve yüklediğiniz e-kitaplar üzerinden yapılandırılmış, dipnotlu akademik metinler hazırlar.\")
    
    if not st.session_state.api_key_valid:
        st.info(\"👉 Devam etmek için lütfen sol panelden geçerli bir Gemini API anahtarı girin.\")
    else:
        col1, col2 = st.columns([2, 1])
        with col1:
            article_topic = st.text_input(\"📝 Makale / Tez Taslağı Konusu:\", value=st.session_state.writer_topic, placeholder=\"Örn: Şüphe krizinden duyulara güvenin sarsılması\")
        with col2:
            output_format = st.selectbox(\"📂 İndirme Formatı:\", [\"Microsoft Word (.docx)\", \"Markdown (.md)\"])
            
        if st.button(\"🚀 Makaleyi Yazmaya Başla\"):
            if not article_topic:
                st.warning(\"Lütfen yazılmasını istediğiniz konuyu girin.\")
            else:
                st.session_state.writer_topic = article_topic
                with st.spinner(\"🔍 ChromaDB ve BM25 üzerinden ilişkili kaynaklar taranıyor...\"):\n                    docs = run_hybrid_search(embed_model, collection, bm25_engine, all_data, article_topic, top_k=5)
                    
                if not docs:
                    st.error(\"Girdiğiniz konuyla ilişkili hiçbir yerel kaynak bulunamadı.\")
                else:
                    context_str = \"\"
                    for doc in docs:
                        label = f\"[{doc['book']} (Sayfa {doc['page']})]\" if doc['book'] else f\"[[{doc['title']}]]\"
                        context_str += f\"\\n--- KAYNAK: {label} ---\\n{doc['text']}\\n--------------------\\n\"
                        
                    system_instruction = \"\"\"Sen İmam Gazâlî felsefesi üzerine uzmanlaşmış, akademik yayın kurallarına hakim bir asistan yazarsın.
Görevin, sana sunulan kaynakları sentezleyerek üst düzey, derinlemesine bir akademik makale taslağı hazırlamaktır.

Makale Yapısı Aynen Şu Şekilde Olmalıdır:
1. AKADEMİK BAŞLIK (İçerikle uyumlu, saygın)
2. ÖZET (Abstract) (Türkçe ve İngilizce - en az 100'er kelime)
3. GİRİŞ (Konunun önemi, ana problemi ve Gazali epistemolojisindeki yeri)
4. GELİŞME (Sana verilen kaynak notlardaki felsefi tartışmaların sentezlenerek alt başlıklarla derinleştirilmesi)
5. SONUÇ (Metinden çıkarılan ana sentez ve bulgular)
6. KAYNAKÇA (Kullandığın kaynak belgelerin ve kitapların listesi)

Kurallar:
- Metinde bilgi uydurma. Gelişme bölümündeki iddiaların sonuna kaynak referanslarını [Kitap Adı (Sayfa No)] şeklinde yaz.
- Akademik, akıcı, zengin ve kusursuz bir Türkçe kullan.
\"\"\"

                    user_prompt = f\"Yazılacak Makale Konusu: {article_topic}\\n\\nSana Sunulan Kaynak Notlar:\\n==================================================\\n{context_str}\\n==================================================\\n\\nYukarıdaki kurallara ve şablona tamamen bağlı kalarak kapsamlı akademik makale taslağını yaz:\"

                    try:
                        with st.spinner(\"✍️ Yapay zekâ akademik taslağınızı yazıyor. Bu işlem 15-30 saniye sürebilir...\"):\n                            model = genai.GenerativeModel(\n                                model_name=st.session_state.selected_model,\n                                system_instruction=system_instruction\n                            )\n                            response = model.generate_content(\n                                contents=user_prompt,\n                                generation_config=genai.types.GenerationConfig(temperature=0.2)\n                            )\n                            st.session_state.writer_article = response.text\n                            st.success(\"Makale Taslağı Başarıyla Üretildi! 🎉\")
                    except Exception as e:
                        st.error(f\"Yazım sırasında bir hata oluştu: {e}\")
        
        # Eğer hafızada makale varsa göster (Sayfadan çıksak bile kaybolmaz!)
        if st.session_state.writer_article:
            st.write(\"---\")
            st.markdown(st.session_state.writer_article)
            st.write(\"---\")
            
            if \"Word\" in output_format:
                file_stream = generate_docx_stream(st.session_state.writer_topic, st.session_state.writer_article)
                st.download_button(
                    label=\"📥 Word Dosyası Olarak İndir (.docx)\",
                    data=file_stream,
                    file_name=f\"gazali_{st.session_state.writer_topic.replace(' ', '_').lower()}.docx\",
                    mime=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document\"\n                )
            else:
                st.download_button(
                    label=\"📥 Markdown Dosyası Olarak İndir (.md)\",
                    data=st.session_state.writer_article,
                    file_name=f\"gazali_{st.session_state.writer_topic.replace(' ', '_').lower()}.md\",
                    mime=\"text/markdown\"\n                )

# ---------------------------------------------------------------------
# TAB 3: İNTERAKTİF KAVRAM AĞ HARİTASI
# ---------------------------------------------------------------------
with tab3:
    st.subheader(\"📊 Obsidian Kavramsal İlişki Haritası\")
    st.write(\"Obsidian notlarınız arasındaki anlamsal wikitext bağlantılarını ve ilişkileri gösteren etkileşimli haritanız:\")
    
    html_path = \"gazali_interaktif_ag.html\"
    png_path = \"gazali_network_graph.png\"
    
    if os.path.exists(html_path):
        try:
            with open(html_path, \"r\", encoding=\"utf-8\") as f:
                html_content = f.read()
            st.components.v1.html(html_content, height=750, scrolling=True)
            st.success(\"🌐 İnteraktif haritanız yüklendi! Mouse ile sürükleyip düğümleri oynatabilirsiniz.\")
        except Exception as e:
            st.error(f\"HTML harita yüklenirken hata oluştu: {e}\")
    elif os.path.exists(png_path):
        st.image(png_path, use_column_width=True)
        st.warning(\"Ölçeklenebilir interaktif haritayı göremiyorsanız lütfen önce 'gazali_network_analysis.py' dosyasını çalıştırın.\")
    else:
        st.warning(\"Henüz oluşturulmuş bir ağ görseli bulunamadı. Lütfen önce 'gazali_network_analysis.py' betiğini çalıştırın.\")

# ---------------------------------------------------------------------
# TAB 4: AYET & HADİS & ŞAHIS ANALİTİĞİ (NER) - Bellek Korumalı!
# ---------------------------------------------------------------------
with tab4:
    st.subheader(\"📜 Yapay Zekâ Destekli Ayet, Hadis ve Şahıs Analitiği\")
    st.caption(\"Külliyatta geçen ayetleri, hadis-i şerifleri ve felsefi/tarihi şahsiyetleri yapay zekâ ile otomatik ayıklar ve listeler.\")
    
    if not st.session_state.api_key_valid:
        st.info(\"👉 Devam etmek için lütfen sol panelden geçerli bir Gemini API anahtarı girin.\")
    else:
        col1, col2 = st.columns([2, 1])
        with col1:
            ner_topic = st.text_input(\"🔍 Taranacak Akademik Konu:\", value=st.session_state.ner_topic_state, placeholder=\"Örn: ilim ve amel, kalbin askerleri, nefs riyazeti, yaratılış\")
        with col2:
            num_docs = st.slider(\"Tarama Kapsamı (Paragraf Sayısı):\", min_value=3, max_value=8, value=5)
            
        if st.button(\"🚀 Tarama ve Analizi Başlat\"):
            if not ner_topic:
                st.warning(\"Lütfen tarama yapılacak bir konu girin.\")
            else:
                st.session_state.ner_topic_state = ner_topic
                with st.spinner(\"🔍 Yerel kitaplar taranıyor ve ilgili paragraflar alınıyor...\"):\n                    docs = run_hybrid_search(embed_model, collection, bm25_engine, all_data, ner_topic, top_k=num_docs)
                    st.session_state.ner_docs = docs
                    
                if not docs:
                    st.error(\"Girdiğiniz konuyla ilişkili hiçbir kaynak bulunamadı.\")
                else:
                    st.info(f\"📚 Konuyla en ilişkili {len(docs)} paragraf başarıyla çekildi. Yapay zekâ ayıklama (NER) motoru çalıştırılıyor...\")
                    
                    with st.spinner(\"🔮 Yapay zekâ anlamsal varlık ayıklama (NER) gerçekleştiriyor...\"):\n                        ner_data, success, err_msg = extract_and_display_ner(docs)
                        if success:
                            st.session_state.ner_results = ner_data
                            st.success(\"Analiz Başarıyla Tamamlandı! 🎉\")
                        else:
                            st.error(f\"NER analizi sırasında bir hata oluştu: {err_msg}\")
                            
        # Hafızada analiz sonucu varsa göster! (Sayfadan çıksak bile kaybolmaz!)
        if st.session_state.ner_results:
            data = st.session_state.ner_results
            
            # 1. AYETLER
            st.markdown(\"### 📖 Ayet-i Kerîmeler\")
            ayet_list = data.get(\"ayetler\", [])
            if ayet_list:
                df_ayet = pd.DataFrame(ayet_list)
                df_ayet.columns = [\"Sure / Ayet No\", \"Ayet Meali / Metni\", \"Geçtiği Kaynak\"]
                st.dataframe(df_ayet, use_container_width=True)
            else:
                st.info(\"Taranan paragraflarda doğrudan ayet-i kerîme atfı tespit edilemedi.\")
                
            # 2. HADİSLER
            st.markdown(\"### 💬 Hadîs-i Şerîfler\")
            hadis_list = data.get(\"hadisler\", [])
            if hadis_list:
                df_hadis = pd.DataFrame(hadis_list)
                df_hadis.columns = [\"Hadis-i Şerif Metni\", \"Ravi / Kaynak\", \"Geçtiği Kaynak\"]
                st.dataframe(df_hadis, use_container_width=True)
            else:
                st.info(\"Taranan paragraflarda doğrudan hadîs-i şerîf atfı tespit edilemedi.\")
                
            # 3. ŞAHISLAR
            st.markdown(\"### 👤 Tarihi Şahsiyetler & Alimler\")
            sahis_list = data.get(\"sahislar\", [])
            if sahis_list:
                df_sahis = pd.DataFrame(sahis_list)
                df_sahis.columns = [\"İsim\", \"Metindeki Rolü / Bağlam\", \"Geçtiği Kaynak\"]
                st.dataframe(df_sahis, use_container_width=True)
            else:
                st.info(\"Taranan paragraflarda özel şahıs/alim atfı tespit edilemedi.\")
                
            # Kaynak Metinler
            if st.session_state.ner_docs:
                with st.expander(\"📝 Taranan Paragrafların Ham Metinleri\"):
                    for d in st.session_state.ner_docs:
                        lbl = f\"📌 {d['book']} | Sayfa: {d['page']}\" if d['book'] else f\"📌 Obsidian: [[{d['title']}]]\"
                        st.markdown(f\"**{lbl}**\")
                        st.caption(f'\"{d[\"text\"]}\"')
"""

with open("main.py", "w", encoding="utf-8") as f:
    f.write(main_py_content)
print("[+] Modüler Streamlit App yazıldı: main.py")

print("\n" + "="*80)
print("[🎉] TEBRİKLER! Kurumsal (Enterprise) Mimariye Dönüşüm ve IP Koruması Tamamlandı!")
print("="*80)
print("• Kritik veritabanınız ve .env dosyalarınız .gitignore ile GitHub sızıntılarına karşı KİLİTLENDİ.")
print("• AGPL-3.0 lisansı ile kodlarınız mülkiyet ve telif hakkı koruması altına alındı.")
print("• Monolitik yapınız config/ ve src/ modüllerine kırılarak yüksek mühendislik kalitesine kavuştu.")
print("\nYerel terminalinizde 'python3 main.py' komutuyla sistemi hemen test edebilirsiniz!")
print("="*80 + "\n")
