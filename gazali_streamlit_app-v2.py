import os
import re
import sys
import streamlit as st
import chromadb
from sentence_transformers import SentenceTransformer
import google.generativeai as genai
from docx import Document
import io

# =====================================================================
# ISLAMICATE DH - GAZALİ PORTALI: STREAMLIT CO-WRITER & CHATBOT (v1)
# =====================================================================
# Bu web portalı, yerel bilgisayarınızda çalışan ChromaDB + Gemini 3.x
# mimarisini tarayıcınıza taşır. Arayüz üzerinden sohbet edebilir,
# anında akademik makale taslakları üretebilir ve kavram ağınızı izleyebilirsiniz.
# =====================================================================

st.set_page_config(
    page_title="İmam Gazâlî Dijital Beşerî Bilimler Portalı",
    page_icon="🕌",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Dark Theme ve Özel CSS Tasarımı
st.markdown("""
<style>
    .reportview-container {
        background: #121212;
    }
    .main .block-container {
        padding-top: 2rem;
    }
    h1, h2, h3 {
        color: #00adb5 !important;
        font-family: 'Georgia', serif;
    }
    .stButton>button {
        background-color: #00adb5;
        color: white;
        border-radius: 8px;
        border: none;
        padding: 0.5rem 1.5rem;
        font-weight: bold;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        background-color: #007a80;
        border: none;
        box-shadow: 0 4px 12px rgba(0,173,181,0.3);
    }
    .source-box {
        background-color: #1e1e1e;
        border-left: 4px solid #00adb5;
        padding: 12px;
        margin: 8px 0px;
        border-radius: 4px;
    }
</style>
""", unsafe_allow_html=True)

# =====================================================================
# 1. TÜRKÇE BM25 VE HİBRİT ARAMA SINIFI (Arka Uç Motoru)
# =====================================================================
class TurkishBM25:
    """Türkçe uyumlu, hafif ve yerel bir BM25 sınıfı."""
    def __init__(self, corpus, b=0.75, k1=1.5):
        self.b = b
        self.k1 = k1
        self.corpus_size = len(corpus)
        self.avg_dl = 0
        self.doc_lengths = []
        self.doc_freqs = {}
        self.idf = {}
        self.tokenized_corpus = []
        
        self._initialize(corpus)

    def _tokenize(self, text):
        # Türkçe küçük harfe çevirme ve noktalama temizliği
        text = text.replace('I', 'ı').replace('İ', 'i').lower()
        words = re.findall(r'[a-zA-Z0-9çğıöşüçĞIİÖŞÜ]+', text)
        return words

    def _initialize(self, corpus):
        total_length = 0
        for doc in corpus:
            tokens = self._tokenize(doc)
            self.tokenized_corpus.append(tokens)
            self.doc_lengths.append(len(tokens))
            total_length += len(tokens)
            
            # Kelime frekanslarını (Document Frequency) bulalım
            unique_tokens = set(tokens)
            for token in unique_tokens:
                self.doc_freqs[token] = self.doc_freqs.get(token, 0) + 1
                
        self.avg_dl = total_length / self.corpus_size if self.corpus_size > 0 else 1
        
        # IDF değerlerini hesaplayalım
        for token, df in self.doc_freqs.items():
            self.idf[token] = max(0.0001, (self.corpus_size - df + 0.5) / (df + 0.5) + 1)

    def get_scores(self, query):
        query_tokens = self._tokenize(query)
        scores = []
        
        for idx, doc_tokens in enumerate(self.tokenized_corpus):
            score = 0.0
            doc_len = self.doc_lengths[idx]
            
            # Her bir döküman için kelime sayıları
            word_counts = {}
            for token in doc_tokens:
                word_counts[token] = word_counts.get(token, 0) + 1
                
            for q_token in query_tokens:
                if q_token in word_counts:
                    tf = word_counts[q_token]
                    idf_val = self.idf.get(q_token, 0.0)
                    numerator = tf * (self.k1 + 1)
                    denominator = tf + self.k1 * (1 - self.b + self.b * (doc_len / self.avg_dl))
                    score += idf_val * (numerator / denominator)
            scores.append(score)
        return scores

# =====================================================================
# 2. STATEYÖNETİMİ VE INITIALIZATION
# =====================================================================
@st.cache_resource
def load_rag_assets():
    """Vektör tabanını ve E5 modelini önbellekte bir kez yükler."""
    try:
        db_path = "./gazali_chroma_db"
        chroma_client = chromadb.PersistentClient(path=db_path)
        embed_model = SentenceTransformer("intfloat/multilingual-e5-large")
        collection = chroma_client.get_collection(name="gazali_kulliyati")
        
        # Tüm verileri BM25 indeksleme için bir kez çekelim
        all_data = collection.get()
        if not all_data or not all_data["documents"]:
            return None, None, None, None
            
        bm25_engine = TurkishBM25(all_data["documents"])
        return embed_model, collection, bm25_engine, all_data
    except Exception as e:
        st.error(f"Veritabanı yüklenirken hata oluştu: {e}")
        return None, None, None, None

embed_model, collection, bm25_engine, all_data = load_rag_assets()

# Session State'leri Başlat
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "api_key_valid" not in st.session_state:
    st.session_state.api_key_valid = False
if "selected_model" not in st.session_state:
    st.session_state.selected_model = None
if "available_models" not in st.session_state:
    st.session_state.available_models = []

# =====================================================================
# 3. YAN PANEL (SIDEBAR) - GÜVENLİK VE AYARLAR
# =====================================================================
with st.sidebar:
    # Gazali Avatar - Yerel dosya varsa yukle yoksa bizim urettigimiz dunya standardindaki bulut adresinden cek
    avatar_filename = "gazali_avatar.jpg"
    if os.path.exists(avatar_filename):
        st.image(avatar_filename, width=120)
    else:
        st.image("https://lh3.googleusercontent.com/notebooklm/AKYWMX9K4yA8aHDxuZk--tSJQd6l4Yl90LDSO1ZzFe1-lCqrfuHnZ-x8qKQQhdg-ICdLXqiNJcRjAXxq-OY-jWoAIsJCCY6-HApYqxNC9qN-BN3bDYwmO-pkMRoP_sRtSOXJgpjfeiDXisLGUmwi8W4vURHHLtZErQ", width=120)
    st.title("🕌 Gazâlî Portal")
    st.caption("Digital Humanities & RAG Platform v1.0")
    st.write("---")
    
    # API Key Alanı
    api_key_input = st.text_input(
        "🔑 Gemini API Key:",
        type="password",
        value=os.environ.get("GEMINI_API_KEY", ""),
        help="Google AI Studio'dan aldığınız API anahtarı."
    )
    
    if api_key_input:
        os.environ["GEMINI_API_KEY"] = api_key_input
        genai.configure(api_key=api_key_input)
        
        # Canlı Modelleri Listeleme Butonu
        if not st.session_state.api_key_valid:
            try:
                raw_models = genai.list_models()
                models_list = []
                for m in raw_models:
                    if "generateContent" in m.supported_generation_methods:
                        m_name = m.name if m.name.startswith("models/") else f"models/{m.name}"
                        # Eski modelleri ele
                        if not any(x in m_name for x in ["gemini-pro", "gemini-2.5"]):
                            models_list.append(m_name)
                
                # Flash modelleri yukarı sırala
                def model_prio(name):
                    return 1 if "flash" in name.lower() else 5
                models_list.sort(key=model_prio)
                
                st.session_state.available_models = models_list
                st.session_state.selected_model = models_list[0] if models_list else None
                st.session_state.api_key_valid = True
                st.success("API Bağlantısı Başarılı! 🎉")
            except Exception as e:
                st.error(f"Bağlantı Hatası: {e}")
                st.session_state.api_key_valid = False
                
    if st.session_state.api_key_valid:
        st.selectbox(
            "🤖 Aktif Gemini Sürümü:",
            options=st.session_state.available_models,
            index=0,
            key="selected_model_box"
        )
        st.session_state.selected_model = st.session_state.selected_model_box
    
    st.write("---")
    # Veri Tabanı Durumu
    if all_data:
        st.metric(label="📚 Toplam Paragraf Sayısı", value=len(all_data["documents"]))
        st.info("Sisteminizde hem kavram notları hem de orijinal e-kitaplar yüklüdür.")
    else:
        st.warning("Veritabanınız boş! Lütfen önce pipeline veya ingester'ı çalıştırın.")

# =====================================================================
# 4. YARDIMCI METODLAR (Arama ve Doküman Oluşturma)
# =====================================================================
def run_hybrid_search(query, top_k=4):
    """BM25 ve Vektör (E5) sıralamalarını RRF formülüyle sentezler."""
    if not collection or not bm25_engine or not all_data:
        return []
        
    # 1. BM25 Skorlarını Hesapla
    bm25_scores = bm25_engine.get_scores(query)
    bm25_ranked = sorted(range(len(bm25_scores)), key=lambda k: bm25_scores[k], reverse=True)
    
    # 2. Vektör (E5) Skorlarını Hesapla
    formatted_query = f"query: {query}"
    query_vector = embed_model.encode(formatted_query).tolist()
    vector_results = collection.query(
        query_embeddings=[query_vector],
        n_results=len(all_data["documents"])
    )
    
    vector_order = vector_results["ids"][0]
    vector_ranked = [all_data["ids"].index(vid) for vid in vector_order]
    
    # 3. Reciprocal Rank Fusion (RRF) Sentezi
    rrf_scores = {}
    k_constant = 60
    
    # BM25 Ranks
    for rank, idx in enumerate(bm25_ranked):
        rrf_scores[idx] = rrf_scores.get(idx, 0.0) + (1.0 / (k_constant + rank + 1))
        
    # Vector Ranks
    for rank, idx in enumerate(vector_ranked):
        rrf_scores[idx] = rrf_scores.get(idx, 0.0) + (1.0 / (k_constant + rank + 1))
        
    # En iyi sonuçları sırala ve topla
    top_indices = sorted(rrf_scores.keys(), key=lambda x: rrf_scores[x], reverse=True)[:top_k]
    
    final_docs = []
    for idx in top_indices:
        metadata = all_data["metadatas"][idx]
        text = all_data["documents"][idx]
        
        # Vektör ve BM25 sıralarını bulalım
        v_rank = vector_ranked.index(idx) + 1 if idx in vector_ranked else "N/A"
        b_rank = bm25_ranked.index(idx) + 1
        
        final_docs.append({
            "text": text,
            "title": metadata.get("title", "Bilinmeyen"),
            "book": metadata.get("book", ""),
            "page": metadata.get("page", ""),
            "links": metadata.get("links", ""),
            "rrf_score": rrf_scores[idx],
            "v_rank": v_rank,
            "b_rank": b_rank
        })
        
    return final_docs

def generate_docx_stream(title, content):
    """Makaleyi doğrudan indirilebilir Word dosyasına dönüştürür."""
    doc = Document()
    doc.add_heading(title, level=0)
    
    # Paragrafları ayırıp ekleme
    for paragraph in content.split("\n\n"):
        if paragraph.strip():
            if paragraph.startswith("### "):
                doc.add_heading(paragraph.replace("### ", ""), level=2)
            elif paragraph.startswith("## "):
                doc.add_heading(paragraph.replace("## ", ""), level=1)
            else:
                doc.add_paragraph(paragraph)
                
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# =====================================================================
# 5. ANA SAYFA VE TAB TASARIMI
# =====================================================================
st.title("🕌 İmam Gazâlî Akademik Araştırma Portalı")
st.write("Klasik İslâmî teoloji, felsefe ve epistemoloji araştırmalarını yapay zekâ ile buluşturan yerel bilgi yönetim paneli.")

tab1, tab2, tab3 = st.tabs(["💬 Akademik Sohbet (Chatbot)", "✍️ Otomatik Co-Writer", "📊 İnteraktif Kavram Ağ Haritası"])

# ---------------------------------------------------------------------
# TAB 1: AKADEMİK SOHBET (CHATBOT)
# ---------------------------------------------------------------------
with tab1:
    st.subheader("🤖 Hibrit Arama Destekli Soru-Cevap Motoru")
    st.caption("Gelişmiş BM25 kelime araması ve E5-Large anlamsal aramayı birleştirerek sıfır halüsinasyonla çalışır.")
    
    if not st.session_state.api_key_valid:
        st.info("👉 Devam etmek için lütfen sol panelden geçerli bir Gemini API anahtarı girin.")
    else:
        # Eski sohbet geçmişini göster
        for message in st.session_state.chat_history:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
                if "sources" in message and message["sources"]:
                    with st.expander("📚 Grounding (Vektörel & BM25 Kaynak Atıfları)"):
                        for doc in message["sources"]:
                            source_header = f"📌 {doc['book']} | Sayfa: {doc['page']}" if doc['book'] else f"📌 Obsidian Notu: [[{doc['title']}]]"
                            st.markdown(f"**{source_header}** *(RRF Skoru: {doc['rrf_score']:.4f}, Vektör Sırası: {doc['v_rank']}, BM25 Sırası: {doc['b_rank']})*")
                            st.caption(f"\"{doc['text']}\"")
        
        # Kullanıcı Girdisi
        user_query = st.chat_input("İmam Gazali felsefesi veya e-kitaplarınız hakkında sorun...")
        
        if user_query:
            # 1. Kullanıcı mesajını göster ve kaydet
            st.chat_message("user").write(user_query)
            st.session_state.chat_history.append({"role": "user", "content": user_query})
            
            # 2. Arka planda Hibrit Arama yap
            with st.spinner("📚 Yerel arşiviniz taranıyor..."):
                retrieved_docs = run_hybrid_search(user_query, top_k=4)
                
            if not retrieved_docs:
                response_text = "Arşivinizde bu konuyla ilişkili hiçbir belge veya kitap bulunamadı."
                st.chat_message("assistant").write(response_text)
                st.session_state.chat_history.append({"role": "assistant", "content": response_text, "sources": []})
            else:
                # 3. Gemini Prompt İnşa Et
                context_str = ""
                for doc in retrieved_docs:
                    source_label = f"[{doc['book']} (Sayfa {doc['page']})]" if doc['book'] else f"[[{doc['title']}]]"
                    context_str += f"\n--- KAYNAK: {source_label} ---\nİçerik: {doc['text']}\n------------------------\n"
                    
                system_instruction = """Sen İmam Gazali felsefesi ve teolojisi üzerine uzmanlaşmış, akademik dürüstlüğü en üst düzeyde tutan bir yapay zeka asistanısın.
Görevin, kullanıcının sorusuna SADECE sana sunulan 'KAYNAK NOTLAR' kapsamındaki verileri kullanarak yanıt vermektir.

Uyman Gereken Kesin Kurallar:
1. Sana verilen KAYNAK NOTLAR dışındaki hiçbir genel kültür veya internet bilgisini kullanma.
2. Eğer sorunun yanıtı sana sunulan notlarda doğrudan veya dolaylı olarak geçmiyorsa, kesinlikle bilgi UYDURMA (Halüsinasyon üretme).
3. Cevap verirken her cümlenin veya iddiadan sonra hangi belgeden/sayfadan alındığını belirt. (Örn: "...duyuların insanı yanılttığını söyler [[Hissiyyât]] veya [Eyyühel Veled (Sayfa 5)].")
4. Çelişkili veya belirsiz durumlar varsa bunları kendi yorumunu katmadan olduğu gibi aktar.
5. Tamamen Türkçe yanıt ver ve akademik, saygın bir dil kullan.
"""

                user_prompt = f"""Kullanıcı Sorusu: {user_query}

Sana Sunulan Kaynak Notlar:
==================================================
{context_str}
==================================================

Yukarıdaki kaynak notlara ve kurallara tamamen bağlı kalarak soruyu yanıtla:"""

                # 4. Gemini 3.x ile cevap üret
                try:
                    with st.spinner("🔮 Gemini anlamsal sentez gerçekleştiriyor..."):
                        model = genai.GenerativeModel(
                            model_name=st.session_state.selected_model,
                            system_instruction=system_instruction
                        )
                        response = model.generate_content(
                            contents=user_prompt,
                            generation_config=genai.types.GenerationConfig(temperature=0.1)
                        )
                        response_text = response.text
                        
                        # Asistan cevabını göster
                        with st.chat_message("assistant"):
                            st.markdown(response_text)
                            with st.expander("📚 Grounding (Vektörel & BM25 Kaynak Atıfları)"):
                                for doc in retrieved_docs:
                                    source_header = f"📌 {doc['book']} | Sayfa: {doc['page']}" if doc['book'] else f"📌 Obsidian Notu: [[{doc['title']}]]"
                                    st.markdown(f"**{source_header}** *(RRF Skoru: {doc['rrf_score']:.4f}, Vektör Sırası: {doc['v_rank']}, BM25 Sırası: {doc['b_rank']})*")
                                    st.caption(f"\"{doc['text']}\"")
                                    
                        # Geçmişe kaydet
                        st.session_state.chat_history.append({
                            "role": "assistant",
                            "content": response_text,
                            "sources": retrieved_docs
                        })
                except Exception as e:
                    st.error(f"Yapay zeka cevap üretirken hata oluştu: {e}")

# ---------------------------------------------------------------------
# TAB 2: AKADEMİK CO-WRITER (YAZAR)
# ---------------------------------------------------------------------
with tab2:
    st.subheader("✍️ Akıllı Makale ve Tez Taslağı Hazırlayıcı")
    st.caption("Obsidian kütüphaneniz ve yüklediğiniz e-kitaplar üzerinden yapılandırılmış, dipnotlu akademik metinler hazırlar.")
    
    if not st.session_state.api_key_valid:
        st.info("👉 Devam etmek için lütfen sol panelden geçerli bir Gemini API anahtarı girin.")
    else:
        # Form Alanları
        col1, col2 = st.columns([2, 1])
        with col1:
            article_topic = st.text_input("📝 Makale / Tez Taslağı Konusu:", placeholder="Örn: Şüphe krizinden duyulara güvenin sarsılması")
        with col2:
            output_format = st.selectbox("📂 İndirme Formatı:", ["Microsoft Word (.docx)", "Markdown (.md)"])
            
        if st.button("🚀 Makaleyi Yazmaya Başla"):
            if not article_topic:
                st.warning("Lütfen yazılmasını istediğiniz konuyu girin.")
            else:
                with st.spinner("🔍 ChromaDB ve BM25 üzerinden ilişkili kaynaklar taranıyor..."):
                    docs = run_hybrid_search(article_topic, top_k=5)
                    
                if not docs:
                    st.error("Girdiğiniz konuyla ilişkili hiçbir yerel kaynak bulunamadı.")
                else:
                    # Kaynak ve Referans oluşturma
                    context_str = ""
                    sources_list = []
                    for doc in docs:
                        label = f"[{doc['book']} (Sayfa {doc['page']})]" if doc['book'] else f"[[{doc['title']}]]"
                        context_str += f"\n--- KAYNAK: {label} ---\n{doc['text']}\n--------------------\n"
                        sources_list.append(label)
                        
                    system_instruction = """Sen İmam Gazâlî felsefesi üzerine uzmanlaşmış, akademik yayın kurallarına hakim bir asistan yazarsın.
Görevin, sana sunulan kaynakları sentezleyerek üst düzey, derinlemesine bir akademik makale taslağı hazırlamaktır.

Makale Yapısı Aynen Şu Şekilde Olmalıdır:
1. AKADEMİK BAŞLIK (İçerikle uyumlu, saygın)
2. ÖZET (Abstract) (Türkçe ve İngilizce - en az 100'er kelime)
3. GİRİŞ (Konunun önemi, ana problemi ve Gazali epistemolojisindeki yeri)
4. GELİŞME (Sana verilen kaynak notlardaki felsefi tartışmaların sentezlenerek alt başlıklarla derinleştirilmesi)
5. SONUÇ (Metinden çıkarılan ana sentez ve bulgular)
6. KAYNAKÇA (Kullandığın kaynak belgelerin ve kitapların listesi)

Kurallar:
- Makale tamamen sana sunulan kaynak verilere sadık kalarak yazılacaktır. Bilgi uydurma.
- Gelişme bölümündeki her iddianın sonuna, hangi not veya kitaptan alındığını çift köşeli parantez referanslarıyla yaz. (Örn: "...deneyimlerin güvenilmez olduğunu açıklar [Eyyühel Veled (Sayfa 5)] veya [[Yakîn]].")
- Akademik, akıcı, zengin ve kusursuz bir Türkçe kullan.
"""

                    user_prompt = f"""Yazılacak Makale Konusu: {article_topic}

Sana Sunulan Kaynak Notlar:
==================================================
{context_str}
==================================================

Yukarıdaki kurallara ve şablona tamamen bağlı kalarak kapsamlı akademik makale taslağını yaz:"""

                    try:
                        with st.spinner("✍️ Yapay zekâ akademik taslağınızı yazıyor. Bu işlem 15-30 saniye sürebilir..."):
                            model = genai.GenerativeModel(
                                model_name=st.session_state.selected_model,
                                system_instruction=system_instruction
                            )
                            response = model.generate_content(
                                contents=user_prompt,
                                generation_config=genai.types.GenerationConfig(temperature=0.2)
                            )
                            generated_text = response.text
                            
                            st.success("Makale Taslağı Başarıyla Üretildi! 🎉")
                            st.write("---")
                            st.markdown(generated_text)
                            st.write("---")
                            
                            # İndirme Butonu Hazırlama
                            if "Word" in output_format:
                                file_stream = generate_docx_stream(article_topic, generated_text)
                                st.download_button(
                                    label="📥 Word Dosyası Olarak İndir (.docx)",
                                    data=file_stream,
                                    file_name=f"gazali_{article_topic.replace(' ', '_').lower()}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                            else:
                                st.download_button(
                                    label="📥 Markdown Dosyası Olarak İndir (.md)",
                                    data=generated_text,
                                    file_name=f"gazali_{article_topic.replace(' ', '_').lower()}.md",
                                    mime="text/markdown"
                                )
                    except Exception as e:
                        st.error(f"Yazım sırasında bir hata oluştu: {e}")

# ---------------------------------------------------------------------
# TAB 3: İNTERAKTİF KAVRAM AĞ HARİTASI
# ---------------------------------------------------------------------
with tab3:
    st.subheader("📊 Obsidian Kavramsal İlişki Haritası")
    st.write("Obsidian notlarınız arasındaki anlamsal wikitext bağlantılarını ve ilişkileri gösteren etkileşimli haritanız:")
    
    # Yerel dosyaları kontrol et
    html_path = "gazali_interaktif_ag.html"
    png_path = "gazali_network_graph.png"
    
    if os.path.exists(html_path):
        try:
            with open(html_path, "r", encoding="utf-8") as f:
                html_content = f.read()
            # Streamlit iframe içinde HTML göster
            st.components.v1.html(html_content, height=750, scrolling=True)
            st.success("🌐 İnteraktif haritanız yüklendi! Mouse ile sürükleyip düğümleri oynatabilirsiniz.")
        except Exception as e:
            st.error(f"HTML harita yüklenirken hata oluştu: {e}")
    elif os.path.exists(png_path):
        st.image(png_path, use_column_width=True)
        st.warning("Ölçeklenebilir interaktif haritayı göremiyorsanız lütfen önce 'gazali_network_analysis.py' dosyasını çalıştırın.")
    else:
        st.warning("Henüz oluşturulmuş bir ağ görseli bulunamadı. Lütfen önce 'gazali_network_analysis.py' betiğini çalıştırın.")
