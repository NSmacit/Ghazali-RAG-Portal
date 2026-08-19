import io
from docx import Document

def generate_docx_stream(title, content):
    """Makaleyi doğrudan indirilebilir Word dosyasına dönüştürür."""
    doc = Document()
    doc.add_heading(title, level=0)
    
    for paragraph in content.split("\n\n"):
        if paragraph.strip():
            if paragraph.startswith("### "):
                doc.add_heading(paragraph.replace("### ", ""), level=2)
            elif paragraph.startswith("## "):
                doc.add_heading(paragraph.replace("## ", ""), level=1)
            else:
                doc.add_paragraph(paragraph)
                
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
