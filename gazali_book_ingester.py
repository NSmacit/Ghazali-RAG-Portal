import os
import re
import sys

# =====================================================================
# ISLAMICATE DH - GAZALİ BÜYÜK VERİ / KİTAP YÜKLEME ARACI (v1)
# =====================================================================
# Bu program, İmam Gazâlî'nin tam eserlerini (el-Munkız, Tehâfüt, vb.)
# okur, onları felsefi ve anlamsal bütünlüklerini bozmadan paragraflara
# (chunks) böler, vektörleştirir ve ChromaDB yerel veritabanınıza yükler.
#
# Desteklenen formatlar: .txt, .md, .docx, .pdf
# Gerekli Kütüphaneler: pip install chromadb sentence-transformers pypdf python-docx
# =====================================================================

def check_dependencies():
    """Gerekli kütüphanelerin kontrolünü yapar ve eksik olanları bildirir."""
    missing = []
    try:
        import chromadb
    except ImportError:
        missing.append("chromadb")
    try:
        from sentence_transformers import SentenceTransformer
    except ImportError:
        missing.append("sentence-transformers")
    try:
        import pypdf
    except ImportError:
        missing.append("pypdf")
    try:
        import docx
    except ImportError:
        missing.append("python-docx")

    if missing:
        print("\n" + "="*60)
        print("[!] EKSİK KÜTÜPHANE TESPİT EDİLDİ!")
        print("Büyük kitap yükleme aracı için şu kütüphaneler gereklidir:")
        for m in missing:
            print(f"   • {m}")
        print("\nYüklemek için terminale şu komutu yazıp Enter'a basın:")
        print(f"pip3 install {' '.join(missing)}")
        print("="*60 + "\n")
        sys.exit(1)

# Dosya yüklemeden önce kütüphaneleri kontrol edelim (Geliştirme ortamında pas geçilebilir)
if __name__ == "__main__" and "IPYTHON" not in globals():
    try:
        check_dependencies()
    except SystemExit:
        # Eğer kullanıcının bilgisayarı dışında bir test ortamındaysak devam etmesini sağlayalım
        pass

import chromadb
from sentence_transformers import SentenceTransformer

class GazaliBookIngester:
    def __init__(self, db_path="./gazali_chroma_db"):
        self.db_path = db_path
        self.collection_name = "gazali_kulliyati"
        
        print("\n[*] Yerel ChromaDB veritabanı bağlantısı kuruluyor...")
        self.chroma_client = chromadb.PersistentClient(path=self.db_path)
        
        print("[*] Yerel anlamsal model yükleniyor (intfloat/multilingual-e5-large)...")
        self.embed_model = SentenceTransformer("intfloat/multilingual-e5-large")
        
        # Koleksiyonu al veya oluştur (Kosinüs benzerliği ile)
        self.collection = self.chroma_client.get_or_create_collection(
            name=self.collection_name,
            metadata={"hnsw:space": "cosine"}
        )

    def extract_text(self, file_path):
        """
        Dosya türüne göre (txt, md, docx, pdf) metin içeriğini çeker.
        """
        ext = os.path.splitext(file_path)[1].lower()
        text_content = ""
        
        if ext in [".txt", ".md"]:
            # Türkçe karakter kodlaması uyumluluğu için utf-8 ve windows-1254/iso-8859-9 fallbacks
            encodings = ["utf-8", "windows-1254", "iso-8859-9", "latin-1"]
            for enc in encodings:
                try:
                    with open(file_path, "r", encoding=enc) as f:
                        text_content = f.read()
                    print(f"[+] Metin dosyası başarıyla okundu (Kodlama: {enc})")
                    break
                except UnicodeDecodeError:
                    continue
            if not text_content:
                raise ValueError("Dosya okunamadı! Lütfen UTF-8 kodlamasında olduğundan emin olun.")
                
        elif ext == ".docx":
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text_content = "\n".join(full_text)
            print(f"[+] Word belgesi başarıyla okundu ({len(doc.paragraphs)} paragraf)")
            
        elif ext == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(file_path)
            pages_text = []
            print(f"[*] PDF sayfa sayısı: {len(reader.pages)}")
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    # Sayfa numarası ekleyerek saklayalım
                    pages_text.append(f"\n--- SAYFA {idx+1} ---\n{page_text}")
            text_content = "\n".join(pages_text)
            print(f"[+] PDF belgesi başarıyla okundu ({len(reader.pages)} sayfa)")
            
        else:
            raise ValueError(f"Desteklenmeyen dosya formatı: {ext}. Lütfen .txt, .md, .docx veya .pdf kullanın.")
            
        return text_content.strip()

    def recursive_chunk_text(self, text, max_chunk_size=1200, overlap=200):
        """
        Büyük metinleri felsefi bütünlüğü bozmayacak şekilde akıllıca bölümler.
        Önce paragrafları (\n\n), sonra cümleleri (. ), en son kelimeleri dener.
        """
        # Adım 1: Sayfa numaralarını ve başlıkları koru
        # Paragraf bazlı bölme
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = ""
        
        current_page = "Bilinmiyor"
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
                
            # Eğer sayfa numarası belirteci varsa güncelleyelim
            page_match = re.search(r'--- SAYFA (\d+) ---', para)
            if page_match:
                current_page = page_match.group(1)
                # Belirteci metinden çıkaralım ama sayfa bilgisini metadata'da tutacağız
                para = re.sub(r'--- SAYFA \d+ ---', '', para).strip()
            
            # Eğer ekleyeceğimiz paragraf limiti aşmıyorsa ekleyelim
            if len(current_chunk) + len(para) + 2 <= max_chunk_size:
                current_chunk += ("\n\n" + para if current_chunk else para)
            else:
                # Limit aşılıyorsa mevcut birikmiş chunk'ı kaydet
                if current_chunk:
                    chunks.append({
                        "text": current_chunk,
                        "page": current_page
                    })
                
                # Eğer yeni paragrafın kendisi tek başına limiti aşıyorsa onu cümle cümle bölelim
                if len(para) > max_chunk_size:
                    sentences = re.split(r'(?<=[.!?])\s+', para)
                    sub_chunk = ""
                    for sent in sentences:
                        if len(sub_chunk) + len(sent) + 1 <= max_chunk_size:
                            sub_chunk += (" " + sent if sub_chunk else sent)
                        else:
                            if sub_chunk:
                                chunks.append({
                                    "text": sub_chunk,
                                    "page": current_page
                                })
                            sub_chunk = sent
                    if sub_chunk:
                        current_chunk = sub_chunk
                else:
                    # Yeni paragrafı bir sonraki chunk'ın başlangıcı yapalım (Overlap/örtüşme ekleyerek)
                    # Overlap için bir önceki chunk'ın son kelimelerini alalım
                    overlap_text = ""
                    if current_chunk:
                        words = current_chunk.split()
                        overlap_text = " ".join(words[-overlap//6:]) # Yaklaşık örtüşme kelime sayısı
                    
                    current_chunk = (overlap_text + "\n\n" + para if overlap_text else para)
                    
        # Son kalan chunk'ı da ekle
        if current_chunk:
            chunks.append({
                "text": current_chunk,
                "page": current_page
            })
            
        return chunks

    def ingest_book(self, file_path, book_title=None, max_chunk_size=1200, overlap=200):
        """
        Büyük kitabı okur, anlamsal parçalara böler, vektörleştirir ve ChromaDB'ye yükler.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"[!] Belirtilen kitap dosyası bulunamadı: {file_path}")
            
        file_name = os.path.basename(file_path)
        if not book_title:
            # Dosya adından temiz bir kitap başlığı üretelim (Örn: el-munkiz-mine-d-dalal)
            book_title = os.path.splitext(file_name)[0].replace("_", " ").replace("-", " ").title()
            
        print(f"\n[*] '{book_title}' eseri okunuyor...")
        raw_text = self.extract_text(file_path)
        
        print("[*] Akıllı anlamsal parçalama (Chunking) yapılıyor...")
        chunks_info = self.recursive_chunk_text(raw_text, max_chunk_size, overlap)
        total_chunks = len(chunks_info)
        print(f"[+] Metin {total_chunks} adet anlamsal paragrafa bölündü.")
        
        documents = []
        embeddings = []
        metadatas = []
        ids = []
        
        print("\n[*] Metinler vektörleştiriliyor (Embedding)...")
        print("[i] İlerleme Çubuğu:")
        
        # ChromaDB'ye her kitabın paragraflarını yüklerken benzersiz kimlikler üreteceğiz
        # Format: kitap_adi_chunk_001
        slugified_title = re.sub(r'[^a-zA-Z0-9]', '_', book_title.lower())
        
        for idx, chunk in enumerate(chunks_info):
            text = chunk["text"]
            page = chunk["page"]
            
            # E5 model ön eki (Dökümanlar için passage:)
            formatted_passage = f"passage: {text}"
            
            # Vektör üretimi
            vector = self.embed_model.encode(formatted_passage).tolist()
            
            # Benzersiz ID
            chunk_id = f"book_{slugified_title}_chunk_{idx+1:04d}"
            
            documents.append(text)
            embeddings.append(vector)
            metadatas.append({
                "title": f"{book_title} (Paragraf {idx+1})",
                "book": book_title,
                "page": str(page),
                "chunk_index": idx + 1,
                "source": file_name,
                "links": "" # Büyük kitaplarda başlangıçta boş kalabilir, zamanla linklenebilir
            })
            ids.append(chunk_id)
            
            # Konsolda şık bir ilerleme göstergesi
            progress = (idx + 1) / total_chunks
            bar_length = 30
            filled_length = int(round(bar_length * progress))
            bar = '█' * filled_length + '-' * (bar_length - filled_length)
            sys.stdout.write(f'\r    |{bar}| {progress*100:.1f}% ({idx+1}/{total_chunks} Paragraf)')
            sys.stdout.flush()
            
        print("\n\n[*] Veriler ChromaDB'ye toplu olarak yazılıyor (Upsert)...")
        
        # ChromaDB'nin tek seferde çok büyük verileri yazarken hata vermemesi için batch'ler halinde yükleyelim
        batch_size = 100
        for i in range(0, len(ids), batch_size):
            self.collection.upsert(
                ids=ids[i:i+batch_size],
                embeddings=embeddings[i:i+batch_size],
                metadatas=metadatas[i:i+batch_size],
                documents=documents[i:i+batch_size]
            )
            
        print(f"\n[🎉] BAŞARILI! '{book_title}' eseri RAG veritabanına tamamen yüklendi.")
        print(f"    • Eklenen Paragraf Sayısı: {total_chunks}")
        print(f"    • Veritabanı ID Aralığı: book_{slugified_title}_chunk_0001 -> book_{slugified_title}_chunk_{total_chunks:04d}")
        print("="*60)

    def test_search(self, query, top_k=3):
        """Yüklenen yeni verileri test etmek için hızlı bir anlamsal arama yapar."""
        formatted_query = f"query: {query}"
        query_vector = self.embed_model.encode(formatted_query).tolist()
        
        results = self.collection.query(
            query_embeddings=[query_vector],
            n_results=top_k
        )
        
        print(f"\n🔍 YENİ VERİLER ÜZERİNDE DENEME ARAMASI: '{query}'")
        print("-" * 60)
        for i in range(len(results['ids'][0])):
            doc_id = results['ids'][0][i]
            score = results['distances'][0][i]
            similarity = 1 - (score / 2)
            metadata = results['metadatas'][0][i]
            text = results['documents'][0][i]
            
            # Kitap veya normal not kontrolü
            source_info = f"Kitap: {metadata.get('book', 'Bilinmeyen Kitap')}" if 'book' in metadata else f"Obsidian Notu: {metadata.get('title')}"
            page_info = f" | Sayfa: {metadata.get('page')}" if 'page' in metadata and metadata.get('page') != 'Bilinmiyor' else ""
            
            print(f"📌 {source_info}{page_info} (Benzerlik: %{similarity*100:.1f})")
            print(f"📖 Metin Alıntısı:\n{text[:200]}...")
            print("-" * 60)


# =====================================================================
# ETKİLEŞİMLİ VERİ YÜKLEME KONSOLU (Interactive Ingester CLI)
# =====================================================================
if __name__ == "__main__":
    print("\n" + "="*60)
    print("📚 İMAM GAZALİ KÜLLİYATI BÜYÜK VERİ / KİTAP YÜKLEME KONSOLU")
    print("="*60)
    
    # 1. Pipeline'ı Başlat
    try:
        ingester = GazaliBookIngester()
        
        print("\n[+] Sistem hazır! Yüklemek istediğiniz kitabın tam yolunu yazın.")
        print("[i] Örn: ~/Desktop/el-munkiz.pdf veya ./Tehafut.docx veya ./el_munkiz.txt")
        print("[i] Çıkmak için 'q' yazabilirsiniz.\n")
        
        while True:
            file_input = input("👉 Yüklemek İstediğiniz Kitabın Dosya Yolu: ").strip()
            
            if not file_input:
                continue
            if file_input.lower() in ['q', 'exit', 'quit']:
                print("\n[*] Oturum kapatılıyor. Veri yükleme başarılı!")
                break
                
            # Yolun başındaki ~ işaretini Mac kullanıcı diziniyle değiştirelim
            file_path = os.path.expanduser(file_input)
            
            if not os.path.exists(file_path):
                print(f"\n[❌] HATA: Dosya belirtilen yolda bulunamadı! Yol: {file_path}")
                print("[i] Lütfen dosyanın adını ve yolunu doğru yazdığınızdan emin olun.\n")
                continue
                
            # Kitap başlığı sorma
            book_title = input("👉 Kitabın Adı (Örnek: el-Munkızu mine'd-Dalâl) \n   (Otomatik belirlemek için Enter'a basın): ").strip()
            if not book_title:
                book_title = None
                
            try:
                # Kitabı yükle
                ingester.ingest_book(file_path, book_title=book_title)
                
                # Başarılı yükleme sonrası hemen test etme teklifi
                test_choice = input("👉 Yüklenen kitap üzerinde test sorgulaması yapmak ister misiniz? (y/n): ").strip().lower()
                if test_choice == 'y':
                    while True:
                        test_query = input("\n🔍 Aramak istediğiniz kavram veya soru (Çıkmak için Enter'a basın): ").strip()
                        if not test_query:
                            break
                        ingester.test_search(test_query)
                print("\n" + "="*50)
                break
                
            except Exception as e:
                print(f"\n[❌] Veri yükleme sırasında bir hata oluştu: {e}\n")
                
    except Exception as e:
        print(f"\n[❌] Sistem başlatılamadı: {e}")
