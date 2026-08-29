import os
import re
import sys
import streamlit as st
import chromadb
from sentence_transformers import SentenceTransformer
import google.generativeai as genai
from docx import Document
import io
import json
import pandas as pd
from gazali_semantic_cache import GazaliSemanticCache

# Tek kaynak: hibrit arama motoru mantigi src/search.py'da yasar (tek noktada bakim).
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from src.search import run_hybrid_search as _src_hybrid_search

# =====================================================================
# ISLAMICATE DH — GAZÂLÎ PORTAL: STREAMLIT CO-WRITER & CHATBOT
# =====================================================================
# Yerel ChromaDB + Gemini mimarisini, Claude tarzı bir sohbet arayuzune
# tasiyan premium portal. Sol panelde bolum navigasyonu ve sohbet gecmisi,
# altta sabit sohbet cubugu; tema .streamlit/config.toml ile yonetilir.
# =====================================================================

st.set_page_config(
    page_title="Gazâlî Portal — Dijital Beşerî Bilimler",
    page_icon="🕌",
    layout="wide",
    initial_sidebar_state="expanded",
)

AVATAR_PATH = "gazali_avatar_v2.jpg"
ASSISTANT_AVATAR = AVATAR_PATH if os.path.exists(AVATAR_PATH) else ":material/auto_awesome:"

# ---------------------------------------------------------------------
# Kullanicinin talebi uzerine: portal ismine yakisir premium dokunuslar.
# Renk/font/tema config.toml'da; burada yalnizca native olarak yapilamayan
# hero baslik (altin gradyan) ve sohbet-gecmisi liste gorunumu var.
# ---------------------------------------------------------------------
st.markdown("""
<style>
    /* Ana icerik genisligini okunur tut */
    .main .block-container { max-width: 1180px; padding-top: 1.4rem; }

    /* Premium hero baslik — altin gradyan, Lora serif */
    .gazali-hero { margin: 0.2rem 0 1.1rem 0; }
    .gazali-hero-title {
        font-family: 'Lora', Georgia, serif;
        font-size: 2.05rem; font-weight: 700; line-height: 1.15;
        margin: 0; letter-spacing: 0.2px;
        background: linear-gradient(100deg, #f4d9a0 0%, #d9b45f 45%, #14b8a6 130%);
        -webkit-background-clip: text; background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    .gazali-hero-sub {
        color: #93a7a0; font-size: 0.98rem; margin: 0.35rem 0 0 0;
        font-weight: 400;
    }
    .gazali-hero-rule {
        height: 2px; width: 74px; margin-top: 0.7rem; border-radius: 2px;
        background: linear-gradient(90deg, #d9b45f, rgba(20,184,166,0.15));
    }

    /* Marka blogu (sidebar ustu) */
    .gazali-brand { text-align:center; padding: 0.1rem 0 0.4rem 0; }
    .gazali-brand h2 {
        font-family:'Lora',serif; font-size:1.28rem; margin:0.35rem 0 0 0;
        color:#e9d8ac;
    }
    .gazali-brand p { color:#7f938c; font-size:0.74rem; margin:0.1rem 0 0 0; letter-spacing:0.4px; }

    /* Sohbet gecmisi: tertiary butonlari sola yasli, tek satir, elipsli liste ogesi yap */
    section[data-testid="stSidebar"] div[data-testid="stButton"] > button[kind="tertiary"] {
        text-align: left; justify-content: flex-start; width: 100%;
        padding: 0.32rem 0.55rem; color: #b9c7c1; font-weight: 400;
        white-space: nowrap; overflow: hidden; text-overflow: ellipsis; display: block;
    }
    section[data-testid="stSidebar"] div[data-testid="stButton"] > button[kind="tertiary"]:hover {
        background: rgba(20,184,166,0.10); color: #e8efec;
    }
    /* Aktif sohbet (primary) — sol altin serit hissi */
    section[data-testid="stSidebar"] div[data-testid="stButton"] > button[kind="primary"] {
        text-align: left; justify-content: flex-start; width: 100%;
        padding: 0.32rem 0.55rem; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; display: block;
    }
    .sidebar-section-label {
        color:#6f847d; font-size:0.72rem; text-transform:uppercase; letter-spacing:1px;
        margin:0.9rem 0 0.2rem 0.2rem; font-weight:600;
    }
</style>
""", unsafe_allow_html=True)


# =====================================================================
# 1. TÜRKÇE BM25 (Arka Uç Motoru — load_rag_assets tarafindan kullanilir)
# =====================================================================
class TurkishBM25:
    """Türkçe uyumlu, hafif ve yerel bir BM25 sınıfı."""
    def __init__(self, corpus, b=0.75, k1=1.5):
        self.b = b
        self.k1 = k1
        self.corpus_size = len(corpus)
        self.avg_dl = 0
        self.doc_lengths = []
        self.doc_freqs = {}
        self.idf = {}
        self.tokenized_corpus = []
        self._initialize(corpus)

    def _tokenize(self, text):
        text = text.replace('I', 'ı').replace('İ', 'i').lower()
        return re.findall(r'[a-zA-Z0-9çğıöşüçĞIİÖŞÜ]+', text)

    def _initialize(self, corpus):
        total_length = 0
        for doc in corpus:
            tokens = self._tokenize(doc)
            self.tokenized_corpus.append(tokens)
            self.doc_lengths.append(len(tokens))
            total_length += len(tokens)
            for token in set(tokens):
                self.doc_freqs[token] = self.doc_freqs.get(token, 0) + 1
        self.avg_dl = total_length / self.corpus_size if self.corpus_size > 0 else 1
        for token, df in self.doc_freqs.items():
            self.idf[token] = max(0.0001, (self.corpus_size - df + 0.5) / (df + 0.5) + 1)

    def get_scores(self, query):
        query_tokens = self._tokenize(query)
        scores = []
        for idx, doc_tokens in enumerate(self.tokenized_corpus):
            score = 0.0
            doc_len = self.doc_lengths[idx]
            word_counts = {}
            for token in doc_tokens:
                word_counts[token] = word_counts.get(token, 0) + 1
            for q_token in query_tokens:
                if q_token in word_counts:
                    tf = word_counts[q_token]
                    idf_val = self.idf.get(q_token, 0.0)
                    numerator = tf * (self.k1 + 1)
                    denominator = tf + self.k1 * (1 - self.b + self.b * (doc_len / self.avg_dl))
                    score += idf_val * (numerator / denominator)
            scores.append(score)
        return scores


# =====================================================================
# 2. VARLIKLARIN YÜKLENMESİ (önbellekli)
# =====================================================================
@st.cache_resource
def load_rag_assets():
    """Vektör tabanını ve E5 modelini önbellekte bir kez yükler."""
    try:
        chroma_client = chromadb.PersistentClient(path="./gazali_chroma_db")
        embed_model = SentenceTransformer("intfloat/multilingual-e5-large")
        collection = chroma_client.get_collection(name="gazali_kulliyati")
        all_data = collection.get()
        if not all_data or not all_data["documents"]:
            return None, None, None, None
        bm25_engine = TurkishBM25(all_data["documents"])
        return embed_model, collection, bm25_engine, all_data
    except Exception as e:
        st.error(f"Veritabanı yüklenirken hata oluştu: {e}")
        return None, None, None, None


embed_model, collection, bm25_engine, all_data = load_rag_assets()


def run_hybrid_search(query, top_k=4):
    """Tek kaynak hibrit arama motoruna (src/search.py) delege eder."""
    return _src_hybrid_search(embed_model, collection, bm25_engine, all_data, query, top_k)


def generate_docx_stream(title, content):
    """Makaleyi doğrudan indirilebilir Word dosyasına dönüştürür."""
    doc = Document()
    doc.add_heading(title, level=0)
    for paragraph in content.split("\n\n"):
        if paragraph.strip():
            if paragraph.startswith("### "):
                doc.add_heading(paragraph.replace("### ", ""), level=2)
            elif paragraph.startswith("## "):
                doc.add_heading(paragraph.replace("## ", ""), level=1)
            else:
                doc.add_paragraph(paragraph)
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# =====================================================================
# 3. OTURUM DURUMU (Session State) — konusma gecmisi dahil
# =====================================================================
def _new_conversation():
    st.session_state.conv_counter += 1
    cid = st.session_state.conv_counter
    st.session_state.conversations[cid] = {"title": "Yeni sohbet", "messages": []}
    st.session_state.active_conv = cid
    return cid


def _delete_conversation(cid):
    """Bir sohbeti siler; aktif sohbet silindiyse başka birine (yoksa yeni bir sohbete) geçer."""
    st.session_state.conversations.pop(cid, None)
    if st.session_state.active_conv == cid:
        if st.session_state.conversations:
            st.session_state.active_conv = next(reversed(st.session_state.conversations))
        else:
            _new_conversation()


def _init_state():
    st.session_state.setdefault("conversations", {})
    st.session_state.setdefault("active_conv", None)
    st.session_state.setdefault("conv_counter", 0)
    st.session_state.setdefault("api_key_valid", False)
    st.session_state.setdefault("selected_model", None)
    st.session_state.setdefault("available_models", [])
    st.session_state.setdefault("co_writer_result", None)
    st.session_state.setdefault("co_writer_topic_saved", None)
    st.session_state.setdefault("co_writer_format_saved", None)
    st.session_state.setdefault("ner_results", None)
    st.session_state.setdefault("ner_raw_docs", None)
    st.session_state.setdefault("ner_topic_saved", None)
    if not st.session_state.conversations:
        _new_conversation()
    if st.session_state.active_conv not in st.session_state.conversations:
        st.session_state.active_conv = next(iter(st.session_state.conversations))


_init_state()


def _hero(title, subtitle):
    st.markdown(
        f"<div class='gazali-hero'><h1 class='gazali-hero-title'>{title}</h1>"
        f"<p class='gazali-hero-sub'>{subtitle}</p><div class='gazali-hero-rule'></div></div>",
        unsafe_allow_html=True,
    )


def _api_gate():
    """API anahtari yoksa nazik bir uyari gosterir ve True dondurur (durdur)."""
    if not st.session_state.api_key_valid:
        st.info("Bu bölüm yapay zekâ kullanır (premium). Sol panelden bir Gemini API anahtarı girerek "
                "deneyebilir ya da **Keşfet & Ara** sekmesinden kaynakları ücretsiz araştırabilirsiniz.",
                icon=":material/lock:")
        return True
    return False


def _show_llm_error(e, context=""):
    """Gemini hatalarını kullanıcı dostu biçimde gösterir.
    Kota/oran limiti (429) durumunda ham hata yerine nazik bir yönlendirme sunar."""
    msg = str(e)
    low = msg.lower()
    is_quota = any(t in low for t in [
        "429", "quota", "exceeded", "resource_exhausted", "resourceexhausted",
        "rate limit", "rate-limit", "too many requests",
    ])
    if is_quota:
        st.warning(
            "Ücretsiz Gemini kotası şu an dolu görünüyor. Şunları deneyebilirsiniz:\n\n"
            "• Sol panelden **başka bir Gemini modeli** seçin (her modelin ayrı günlük kotası vardır).\n"
            "• Birkaç dakika bekleyip **tekrar deneyin** (dakikalık limit kısa sürede yenilenir).\n"
            "• Günlük ücretsiz kota Pasifik saatiyle gece yarısı sıfırlanır.",
            icon=":material/hourglass_top:",
        )
    else:
        st.error(f"{context}{msg}", icon=":material/error:")


def _render_sources(docs, label="Kaynak atıfları (Vektör + BM25)"):
    with st.expander(label, icon=":material/menu_book:"):
        for doc in docs:
            if doc.get("book"):
                header = f"**{doc['book']}** · Sayfa {doc['page']}"
            else:
                header = f"**Obsidian notu:** [[{doc['title']}]]"
            st.markdown(
                f":material/bookmark: {header}  \n"
                f":gray[RRF {doc['rrf_score']:.4f} · Vektör #{doc['v_rank']} · BM25 #{doc['b_rank']}]"
            )
            st.caption(f"“{doc['text']}”")


# =====================================================================
# 4. SAYFA: KEŞFET & ARA (Ücretsiz — LLM yok, API anahtarı gerekmez)
# =====================================================================
def page_search():
    _hero("Keşfet & Ara", "Gazâlî külliyatını kaynağıyla ara — kayıt ve API anahtarı gerekmez, tamamen ücretsiz.")

    q = st.text_input(
        "Ara", placeholder="Örn: kalbin hakikati, yakîn, ilim ve amel, filozofların tutarsızlığı…",
        label_visibility="collapsed", icon=":material/search:",
    )

    # Bos durumda ilham verici konu cipleri
    if not q:
        picked = st.pills(
            "Konular",
            ["nefsin bilinmesi", "yakîn ve şüphe", "ilim ve amel",
             "kalbin hakikati", "riyâzet ve nefs terbiyesi", "tevekkül"],
            label_visibility="collapsed", key="search_topics",
        )
        if picked:
            q = picked

    if not q:
        st.caption("Bir kavram veya soru yazın; sistem en ilgili paragrafları eser ve sayfa künyesiyle listeler.")
        return

    with st.spinner("Külliyat taranıyor…"):
        docs = run_hybrid_search(q, top_k=8)

    if not docs:
        st.info("Bu sorguyla ilişkili paragraf bulunamadı.", icon=":material/search_off:")
        return

    st.caption(f"{len(docs)} ilgili paragraf")
    for d in docs:
        with st.container(border=True):
            src = d["book"] if d["book"] else f"Obsidian notu: [[{d['title']}]]"
            page = f" · Sayfa {d['page']}" if d.get("page") and d["page"] != "Bilinmiyor" else ""
            st.markdown(f":material/bookmark: **{src}**{page}")
            st.write(d["text"])
            st.caption(f":gray[RRF {d['rrf_score']:.4f} · Vektör #{d['v_rank']} · BM25 #{d['b_rank']}]")


# =====================================================================
# 5. SAYFA: AKADEMİK SOHBET (CHATBOT)
# =====================================================================
SUGGESTIONS = {
    ":material/psychology: Şüphe krizi ve kesin bilgi (yakîn) arayışı":
        "Gazâlî'nin şüphe krizinden kesin bilgiye (yakîn) ulaşma sürecini anlat.",
    ":material/favorite: Kalbin hakikati ve nefsin bilinmesi":
        "Kalbin hakikati ve marifetü'n-nefs nedir?",
    ":material/balance: Filozofların tutarsızlığı ve metafizik eleştirisi":
        "Gazâlî filozofların metafizik iddialarını neden tutarsız bulur?",
    ":material/school: İlim ve amel ilişkisi (Eyyühe'l-Veled)":
        "Eyyühe'l-Veled'de ilim ve amel ilişkisi nasıl açıklanır?",
}


def _answer_chat(user_query, messages, conv):
    """RAG + semantik önbellek + Gemini akışı; mesajları aktif konuşmaya işler."""
    was_empty = len(messages) == 0
    messages.append({"role": "user", "content": user_query})
    if was_empty:
        conv["title"] = (user_query[:38] + "…") if len(user_query) > 38 else user_query

    st.chat_message("user").markdown(user_query)

    query_vector = embed_model.encode(f"query: {user_query}").tolist()
    cached_response, similarity = st.session_state.semantic_cache.check_cache(user_query, query_vector)

    if cached_response:
        with st.chat_message("assistant", avatar=ASSISTANT_AVATAR):
            st.caption(f":material/bolt: Semantik önbellekten yanıtlandı · %{similarity*100:.1f} benzerlik")
            st.markdown(cached_response)
        messages.append({
            "role": "assistant",
            "content": f":material/bolt: *Semantik önbellekten (%{similarity*100:.1f} benzerlik)*\n\n" + cached_response,
            "sources": [],
        })
        return

    with st.spinner("Yerel arşiviniz taranıyor…"):
        docs = run_hybrid_search(user_query, top_k=4)

    if not docs:
        text = "Arşivinizde bu konuyla ilişkili bir belge veya kitap bulunamadı."
        st.chat_message("assistant", avatar=ASSISTANT_AVATAR).markdown(text)
        messages.append({"role": "assistant", "content": text, "sources": []})
        return

    context_str = ""
    for doc in docs:
        label = f"[{doc['book']} (Sayfa {doc['page']})]" if doc["book"] else f"[[{doc['title']}]]"
        context_str += f"\n--- KAYNAK: {label} ---\nİçerik: {doc['text']}\n------------------------\n"

    system_instruction = """Sen İmam Gazali felsefesi ve teolojisi üzerine uzmanlaşmış, akademik dürüstlüğü en üst düzeyde tutan bir yapay zeka asistanısın.
Görevin, kullanıcının sorusuna SADECE sana sunulan 'KAYNAK NOTLAR' kapsamındaki verileri kullanarak yanıt vermektir.

Uyman Gereken Kesin Kurallar:
1. Sana verilen KAYNAK NOTLAR dışındaki hiçbir genel kültür veya internet bilgisini kullanma.
2. Eğer sorunun yanıtı sana sunulan notlarda doğrudan veya dolaylı olarak geçmiyorsa, kesinlikle bilgi UYDURMA (Halüsinasyon üretme).
3. Cevap verirken her cümlenin veya iddiadan sonra hangi belgeden/sayfadan alındığını belirt. (Örn: "...duyuların insanı yanılttığını söyler [[Hissiyyât]] veya [Eyyühel Veled (Sayfa 5)].")
4. Çelişkili veya belirsiz durumlar varsa bunları kendi yorumunu katmadan olduğu gibi aktar.
5. Tamamen Türkçe yanıt ver ve akademik, saygın bir dil kullan.
"""
    user_prompt = f"""Kullanıcı Sorusu: {user_query}

Sana Sunulan Kaynak Notlar:
==================================================
{context_str}
==================================================

Yukarıdaki kaynak notlara ve kurallara tamamen bağlı kalarak soruyu yanıtla:"""

    try:
        with st.spinner("Gemini anlamsal sentez gerçekleştiriyor…"):
            model = genai.GenerativeModel(
                model_name=st.session_state.selected_model,
                system_instruction=system_instruction,
            )
            response = model.generate_content(
                contents=user_prompt,
                generation_config=genai.types.GenerationConfig(temperature=0.1),
            )
            response_text = response.text
            st.session_state.semantic_cache.set_cache(user_query, response_text, query_vector)

        with st.chat_message("assistant", avatar=ASSISTANT_AVATAR):
            st.markdown(response_text)
            _render_sources(docs)
        messages.append({"role": "assistant", "content": response_text, "sources": docs})
    except Exception as e:
        _show_llm_error(e, "Yanıt üretilirken bir hata oluştu: ")


def page_chat():
    _hero("Akademik Sohbet", "Hibrit arama destekli, kaynak-atıflı soru–cevap motoru.")
    if _api_gate():
        return

    if "semantic_cache" not in st.session_state:
        st.session_state.semantic_cache = GazaliSemanticCache(threshold=0.94)

    conv = st.session_state.conversations[st.session_state.active_conv]
    messages = conv["messages"]

    # Gecmis mesajlar (yukari kayar)
    for message in messages:
        avatar = ASSISTANT_AVATAR if message["role"] == "assistant" else None
        with st.chat_message(message["role"], avatar=avatar):
            st.markdown(message["content"])
            if message.get("sources"):
                _render_sources(message["sources"])

    # Bos sohbette karsilama + oneri cipleri
    pill_query = None
    if not messages:
        st.caption("Başlamak için bir örnek soru seçin ya da kendi sorunuzu yazın.")
        selected = st.pills("Örnek sorular", list(SUGGESTIONS.keys()),
                            label_visibility="collapsed", key="chat_suggestions")
        if selected:
            pill_query = SUGGESTIONS[selected]

    # Alt sabit sohbet cubugu
    typed = st.chat_input("İmam Gazâlî felsefesi veya e-kitaplarınız hakkında sorun…")
    user_query = typed or pill_query
    if user_query:
        _answer_chat(user_query, messages, conv)
        if pill_query and not typed:
            st.rerun()  # karsilama cipini temizlemek icin yeniden ciz


# =====================================================================
# 5. SAYFA: AKADEMİK CO-WRITER
# =====================================================================
def page_cowriter():
    _hero("Akademik Co-Writer", "Külliyatınızdan yapılandırılmış, dipnotlu makale taslakları üretir.")
    if _api_gate():
        return

    c1, c2 = st.columns([2, 1])
    with c1:
        article_topic = st.text_input("Makale / tez taslağı konusu",
                                       placeholder="Örn: Şüphe krizinden duyulara güvenin sarsılması")
    with c2:
        output_format = st.selectbox("İndirme formatı", ["Microsoft Word (.docx)", "Markdown (.md)"])

    if st.button("Makaleyi yazmaya başla", icon=":material/edit_note:", type="primary"):
        if not article_topic:
            st.warning("Lütfen yazılmasını istediğiniz konuyu girin.", icon=":material/warning:")
        else:
            with st.spinner("İlişkili kaynaklar taranıyor…"):
                docs = run_hybrid_search(article_topic, top_k=5)
            if not docs:
                st.error("Girdiğiniz konuyla ilişkili yerel kaynak bulunamadı.", icon=":material/error:")
            else:
                context_str = ""
                for doc in docs:
                    label = f"[{doc['book']} (Sayfa {doc['page']})]" if doc["book"] else f"[[{doc['title']}]]"
                    context_str += f"\n--- KAYNAK: {label} ---\n{doc['text']}\n--------------------\n"

                system_instruction = """Sen İmam Gazâlî felsefesi üzerine uzmanlaşmış, akademik yayın kurallarına hakim bir asistan yazarsın.
Görevin, sana sunulan kaynakları sentezleyerek üst düzey, derinlemesine bir akademik makale taslağı hazırlamaktır.

Makale Yapısı Aynen Şu Şekilde Olmalıdır:
1. AKADEMİK BAŞLIK (İçerikle uyumlu, saygın)
2. ÖZET (Abstract) (Türkçe ve İngilizce - en az 100'er kelime)
3. GİRİŞ (Konunun önemi, ana problemi ve Gazali epistemolojisindeki yeri)
4. GELİŞME (Sana verilen kaynak notlardaki felsefi tartışmaların sentezlenerek alt başlıklarla derinleştirilmesi)
5. SONUÇ (Metinden çıkarılan ana sentez ve bulgular)
6. KAYNAKÇA (Kullandığın kaynak belgelerin ve kitapların listesi)

Kurallar:
- Makale tamamen sana sunulan kaynak verilere sadık kalarak yazılacaktır. Bilgi uydurma.
- Gelişme bölümündeki her iddianın sonuna, hangi not veya kitaptan alındığını çift köşeli parantez referanslarıyla yaz. (Örn: "...deneyimlerin güvenilmez olduğunu açıklar [Eyyühel Veled (Sayfa 5)] veya [[Yakîn]].")
- Akademik, akıcı, zengin ve kusursuz bir Türkçe kullan.
"""
                user_prompt = f"""Yazılacak Makale Konusu: {article_topic}

Sana Sunulan Kaynak Notlar:
==================================================
{context_str}
==================================================

Yukarıdaki kurallara ve şablona tamamen bağlı kalarak kapsamlı akademik makale taslağını yaz:"""
                try:
                    with st.spinner("Yapay zekâ akademik taslağınızı yazıyor (15–30 sn)…"):
                        model = genai.GenerativeModel(
                            model_name=st.session_state.selected_model,
                            system_instruction=system_instruction,
                        )
                        response = model.generate_content(
                            contents=user_prompt,
                            generation_config=genai.types.GenerationConfig(temperature=0.2),
                        )
                        st.session_state.co_writer_result = response.text
                        st.session_state.co_writer_topic_saved = article_topic
                        st.session_state.co_writer_format_saved = output_format
                        st.toast("Makale taslağı üretildi!", icon=":material/check_circle:")
                except Exception as e:
                    _show_llm_error(e, "Yazım sırasında bir hata oluştu: ")

    if st.session_state.co_writer_result:
        st.markdown(st.session_state.co_writer_result)
        topic = st.session_state.co_writer_topic_saved
        if "Word" in (st.session_state.co_writer_format_saved or ""):
            st.download_button(
                "Word (.docx) olarak indir", icon=":material/download:",
                data=generate_docx_stream(topic, st.session_state.co_writer_result),
                file_name=f"gazali_{topic.replace(' ', '_').lower()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        else:
            st.download_button(
                "Markdown (.md) olarak indir", icon=":material/download:",
                data=st.session_state.co_writer_result,
                file_name=f"gazali_{topic.replace(' ', '_').lower()}.md",
                mime="text/markdown",
            )


# =====================================================================
# 6. SAYFA: KAVRAM AĞ HARİTASI
# =====================================================================
def _render_html_file(html_content, height=740):
    """JS iceren HTML'i, surumden bagimsiz sekilde gomer."""
    try:
        st.html(html_content, height=height, unsafe_allow_javascript=True)
    except TypeError:
        st.components.v1.html(html_content, height=height, scrolling=True)


def page_network():
    _hero("Kavram Ağ Haritası", "Obsidian notları arasındaki anlamsal wiki-bağlantılarının etkileşimli haritası.")
    html_path = "gazali_interaktif_ag.html"
    png_path = "gazali_network_graph.png"
    if os.path.exists(html_path):
        try:
            with open(html_path, "r", encoding="utf-8") as f:
                _render_html_file(f.read())
            st.caption("Düğümleri sürükleyerek ilişkileri keşfedebilirsiniz.")
        except Exception as e:
            st.error(f"HTML harita yüklenirken hata oluştu: {e}", icon=":material/error:")
    elif os.path.exists(png_path):
        st.image(png_path, width="stretch")
        st.caption("İnteraktif harita için 'gazali_network_analysis.py' betiğini çalıştırın.")
    else:
        st.warning("Henüz oluşturulmuş bir ağ görseli bulunamadı. Önce 'gazali_network_analysis.py' betiğini çalıştırın.",
                   icon=":material/hub:")


# =====================================================================
# 7. SAYFA: AYET & HADİS & ŞAHIS ANALİTİĞİ (NER)
# =====================================================================
def page_analytics():
    _hero("Ayet · Hadis · Şahıs Analitiği", "Külliyattaki ayet, hadis ve tarihî şahsiyetleri yapay zekâ ile ayıklar.")
    if _api_gate():
        return

    c1, c2 = st.columns([2, 1])
    with c1:
        ner_topic = st.text_input("Taranacak akademik konu",
                                  placeholder="Örn: ilim ve amel, kalbin askerleri, nefs riyazeti")
    with c2:
        num_docs = st.slider("Tarama kapsamı (paragraf)", min_value=3, max_value=8, value=5)

    if st.button("Tarama ve analizi başlat", icon=":material/travel_explore:", type="primary"):
        if not ner_topic:
            st.warning("Lütfen tarama yapılacak bir konu girin.", icon=":material/warning:")
        else:
            with st.spinner("Yerel kitaplar taranıyor…"):
                docs = run_hybrid_search(ner_topic, top_k=num_docs)
            if not docs:
                st.error("Girdiğiniz konuyla ilişkili kaynak bulunamadı.", icon=":material/error:")
            else:
                context_str = ""
                for doc in docs:
                    label = f"[{doc['book']} (Sayfa {doc['page']})]" if doc["book"] else f"[[{doc['title']}]]"
                    context_str += f"\n--- KAYNAK: {label} ---\n{doc['text']}\n--------------------\n"

                system_instruction = """Sen klasik İslami metinler ve teoloji üzerinde uzmanlaşmış, son derece hassas bir metin analitiği (Named Entity Recognition) asistanısın.
Görevin, sana sunulan kaynak metinleri analiz ederek içindeki:
1. AYETLERİ (Eğer varsa; sure adı, ayet numarası, Türkçe meali ve geçtiği kaynak ile birlikte)
2. HADİSLERİ (Eğer varsa; hadisin içeriği, ravisi veya kaynağı, geçtiği yer ile birlikte)
3. ŞAHISLARI (Metinde adı geçen tarihi alimler, filozoflar, peygamberler, sahabeler veya şahsiyetler)
bulup ayıklamaktır.

Çıktıyı KESİNLİKLE ama KESİNLİKLE sadece aşağıdaki JSON şablonuna göre üretmelisin. JSON dışında hiçbir giriş, açıklama, markdown kodu veya düz metin yazmamalısın. Çıktı doğrudan geçerli bir JSON olmalıdır:
{
  "ayetler": [
    {"sure_ayet": "Sure Adı ve Ayet No", "metin": "Ayeti kerimenin meali veya içeriği", "kaynak": "Geçtiği kitap/not ve sayfa"}
  ],
  "hadisler": [
    {"metin": "Hadis-i şerifin meali veya içeriği", "ravi_kaynak": "Zikredilen ravi veya hadis kaynağı", "kaynak": "Geçtiği kitap/not ve sayfa"}
  ],
  "sahislar": [
    {"isim": "Kişinin Adı", "rol_baglam": "Metindeki rolü veya hangi bağlamda zikredildiği", "kaynak": "Geçtiği kitap/not ve sayfa"}
  ]
}

Eğer metinde ayet, hadis veya şahıs yoksa, ilgili listeyi boş bırak: []
"""
                user_prompt = f"""Analiz Edilecek Kaynak Metinler:\n==================================================\n{context_str}\n==================================================\n\nYukarıdaki kaynakları tara ve JSON formatında ayıklama sonuçlarını döndür:"""
                try:
                    with st.spinner("Yapay zekâ varlık ayıklama (NER) gerçekleştiriyor…"):
                        model = genai.GenerativeModel(
                            model_name=st.session_state.selected_model,
                            system_instruction=system_instruction,
                        )
                        response = model.generate_content(
                            contents=user_prompt,
                            generation_config=genai.types.GenerationConfig(
                                temperature=0.1, response_mime_type="application/json"
                            ),
                        )
                        raw_json = response.text.strip()
                        if raw_json.startswith("```json"):
                            raw_json = raw_json.split("```json", 1)[1]
                        if raw_json.endswith("```"):
                            raw_json = raw_json.rsplit("```", 1)[0]
                        st.session_state.ner_results = json.loads(raw_json.strip())
                        st.session_state.ner_raw_docs = docs
                        st.session_state.ner_topic_saved = ner_topic
                        st.toast("Analiz tamamlandı!", icon=":material/check_circle:")
                except Exception as e:
                    _show_llm_error(e, "NER analizi sırasında bir hata oluştu: ")

    if st.session_state.ner_results:
        data = st.session_state.ner_results
        docs = st.session_state.ner_raw_docs

        t_ayet, t_hadis, t_sahis = st.tabs([
            ":material/menu_book: Ayet-i kerîmeler",
            ":material/format_quote: Hadîs-i şerîfler",
            ":material/groups: Şahsiyetler",
        ])
        with t_ayet:
            rows = data.get("ayetler", [])
            if rows:
                df = pd.DataFrame(rows)
                df.columns = ["Sure / Ayet No", "Ayet Meali / Metni", "Geçtiği Kaynak"]
                st.dataframe(df, width="stretch", hide_index=True)
            else:
                st.caption("Taranan paragraflarda doğrudan ayet atfı tespit edilemedi.")
        with t_hadis:
            rows = data.get("hadisler", [])
            if rows:
                df = pd.DataFrame(rows)
                df.columns = ["Hadis-i Şerif Metni", "Ravi / Kaynak", "Geçtiği Kaynak"]
                st.dataframe(df, width="stretch", hide_index=True)
            else:
                st.caption("Taranan paragraflarda doğrudan hadîs atfı tespit edilemedi.")
        with t_sahis:
            rows = data.get("sahislar", [])
            if rows:
                df = pd.DataFrame(rows)
                df.columns = ["İsim", "Metindeki Rolü / Bağlam", "Geçtiği Kaynak"]
                st.dataframe(df, width="stretch", hide_index=True)
            else:
                st.caption("Taranan paragraflarda özel şahıs/alim atfı tespit edilemedi.")

        with st.expander("Taranan paragrafların ham metinleri", icon=":material/description:"):
            for d in docs:
                lbl = f"**{d['book']}** · Sayfa {d['page']}" if d["book"] else f"**Obsidian:** [[{d['title']}]]"
                st.markdown(f":material/bookmark: {lbl}")
                st.caption(f"“{d['text']}”")


# =====================================================================
# 8. YAN PANEL (SIDEBAR): marka · sohbet geçmişi · ayarlar
# =====================================================================
if os.path.exists(AVATAR_PATH):
    try:
        st.logo(AVATAR_PATH, size="large")
    except Exception:
        pass

with st.sidebar:
    st.markdown(
        "<div class='gazali-brand'><h2>🕌 Gazâlî Portal</h2>"
        "<p>DİJİTAL BEŞERÎ BİLİMLER & RAG</p></div>",
        unsafe_allow_html=True,
    )

# Bolum navigasyonu (st.navigation menusu sidebar'in ustune yerlesir)
pages = [
    st.Page(page_search, title="Keşfet & Ara", icon=":material/search:", default=True),
    st.Page(page_chat, title="Sohbet", icon=":material/forum:"),
    st.Page(page_cowriter, title="Co-Writer", icon=":material/edit_note:"),
    st.Page(page_network, title="Ağ Haritası", icon=":material/hub:"),
    st.Page(page_analytics, title="Analitik", icon=":material/travel_explore:"),
]
nav = st.navigation(pages, position="sidebar")

with st.sidebar:
    # --- Sohbet gecmisi (Claude tarzi) ---
    st.markdown("<div class='sidebar-section-label'>Sohbet geçmişi</div>", unsafe_allow_html=True)
    if st.button("Yeni sohbet", icon=":material/add:", width="stretch"):
        # Zaten bos bir aktif sohbet varsa yenisini olusturma
        active = st.session_state.conversations.get(st.session_state.active_conv)
        if not active or active["messages"]:
            _new_conversation()
        st.rerun()

    # Sadece mesaj almis sohbetler listelenir (bos aktif sohbet listeyi kirletmez)
    listed = [cid for cid in reversed(list(st.session_state.conversations.keys()))
              if st.session_state.conversations[cid]["messages"]]
    if listed:
        for cid in listed:
            c = st.session_state.conversations[cid]
            is_active = cid == st.session_state.active_conv
            row = st.columns([0.82, 0.18], gap="small", vertical_alignment="center")
            with row[0]:
                if st.button(c["title"] or "Yeni sohbet", key=f"conv_{cid}", width="stretch",
                             type="primary" if is_active else "tertiary"):
                    st.session_state.active_conv = cid
                    st.rerun()
            with row[1]:
                if st.button(":material/delete:", key=f"del_{cid}", type="tertiary",
                             help="Bu sohbeti sil", width="stretch"):
                    _delete_conversation(cid)
                    st.rerun()
    else:
        st.caption("Henüz kayıtlı sohbet yok.")

    # --- Ayarlar & baglanti ---
    st.markdown("<div class='sidebar-section-label'>Ayarlar</div>", unsafe_allow_html=True)
    api_key_input = st.text_input(
        "Gemini API anahtarı", type="password",
        value=os.environ.get("GEMINI_API_KEY", ""),
        help="Google AI Studio'dan aldığınız API anahtarı.",
    )
    if api_key_input:
        os.environ["GEMINI_API_KEY"] = api_key_input
        genai.configure(api_key=api_key_input)
        if not st.session_state.api_key_valid:
            try:
                models_list = []
                for m in genai.list_models():
                    if "generateContent" in m.supported_generation_methods:
                        m_name = m.name if m.name.startswith("models/") else f"models/{m.name}"
                        if not any(x in m_name for x in ["gemini-pro", "gemini-2.5"]):
                            models_list.append(m_name)
                models_list.sort(key=lambda name: 1 if "flash" in name.lower() else 5)
                st.session_state.available_models = models_list
                st.session_state.selected_model = models_list[0] if models_list else None
                st.session_state.api_key_valid = True
                st.toast("API bağlantısı başarılı!", icon=":material/check_circle:")
            except Exception as e:
                st.error(f"Bağlantı hatası: {e}", icon=":material/error:")
                st.session_state.api_key_valid = False

    if st.session_state.api_key_valid:
        st.session_state.selected_model = st.selectbox(
            "Aktif Gemini sürümü", options=st.session_state.available_models, index=0
        )
        st.caption(":material/check_circle: Bağlantı etkin")

    if all_data:
        st.metric("Toplam paragraf", f"{len(all_data['documents']):,}".replace(",", "."))
    else:
        st.warning("Veritabanı boş! Önce ingester'ı çalıştırın.", icon=":material/database:")

# Secili sayfayi calistir
nav.run()
