import os
import sys
import re
import chromadb
from sentence_transformers import SentenceTransformer
import google.generativeai as genai

# =====================================================================
# ISLAMICATE DH - GAZALİ AKADEMİK YAZAR SİNERJİSİ (ACADEMIC CO-WRITER v1)
# =====================================================================
# Bu araç, yerel ChromaDB veritabanınızı ve en güçlü Gemini 3.x modellerini
# kullanarak, Obsidian notlarınızdan beslenen, dipnotlu, akademik formatta 
# (Özet, Giriş, Gelişme, Sonuç ve Kaynakça) tam taslak makaleler üretir.
# Markdown (.md) ve Word (.docx) çıktısı üretebilir.
# =====================================================================

class GazaliAcademicCoWriter:
    def __init__(self, db_path="./gazali_chroma_db"):
        self.db_path = db_path
        self.selected_model = None
        self.available_models_list = []
        
        # 1. Yerel Veri Tabanı Bağlantısı
        if not os.path.exists(self.db_path):
            raise FileNotFoundError(
                f"[!] '{self.db_path}' dizini bulunamadı! Lütfen önce "
                "'obsidian_rag_pipeline.py' dosyasını çalıştırarak veritabanını inşa edin."
            )
            
        print("[*] Yerel ChromaDB veritabanı yükleniyor...")
        self.chroma_client = chromadb.PersistentClient(path=self.db_path)
        
        # 2. Embedding Modelini Yükle
        print("[*] intfloat/multilingual-e5-large anlamsal modeli önbellekten çağrılıyor...")
        self.embed_model = SentenceTransformer("intfloat/multilingual-e5-large")
        
        # 3. Koleksiyonu Al
        try:
            self.collection = self.chroma_client.get_collection(name="gazali_kulliyati")
        except Exception as e:
            raise ValueError(
                f"[!] 'gazali_kulliyati' koleksiyonu bulunamadı! Hata: {e}"
            )

        # 4. Gemini API Yapılandırması (v7 Dinamik Keşif Altyapısı)
        self.setup_gemini_api_safely()

    def setup_gemini_api_safely(self):
        """
        Kullanıcının taze API anahtarını alır ve Google sunucularındaki en güncel
        modelleri (Gemini 3.x vb.) dinamik olarak sorgulayıp eşleştirir.
        """
        api_key = os.environ.get("GEMINI_API_KEY")
        
        if not api_key:
            print("\n" + "="*50)
            print("[!] GEMINI_API_KEY bulunamadı.")
            print("Lütfen Google AI Studio'dan aldığınız GÜNCEL API anahtarını girin:")
            import getpass
            api_key = getpass.getpass("API Key: ").strip()
            if not api_key:
                print("[!] Anahtar girilmedi. Program sonlandırılıyor.")
                sys.exit(1)
            os.environ["GEMINI_API_KEY"] = api_key

        genai.configure(api_key=api_key)
        
        print("\n[*] Google sunucularına bağlanılıyor ve güncel model yetkileriniz taranıyor...")
        try:
            raw_models = genai.list_models()
            self.available_models_list = []
            
            for m in raw_models:
                if "generateContent" in m.supported_generation_methods:
                    m_name = m.name if m.name.startswith("models/") else f"models/{m.name}"
                    # Eski kısıtlı modelleri eliyoruz
                    if any(x in m_name for x in ["gemini-pro", "gemini-2.5", "gemini-1.0-pro"]):
                        continue
                    self.available_models_list.append(m_name)
            
            # 2026 model hiyerarşisi önceliklendirmesi (3.7 -> 3.5 -> 3.1 -> flash-latest)
            def model_priority(name):
                name_lower = name.lower()
                if "3.7-flash" in name_lower: return 1
                if "3.5-flash" in name_lower: return 2
                if "3.1-pro" in name_lower: return 3
                if "flash-latest" in name_lower: return 4
                if "flash" in name_lower: return 5
                return 10
            
            self.available_models_list.sort(key=model_priority)
            
            if self.available_models_list:
                self.selected_model = self.available_models_list[0]
                print(f"[+] API Doğrulandı! Seçilen En Güçlü Yapay Zeka Modeli: {self.selected_model}")
            else:
                self.selected_model = "models/gemini-flash-latest"
                print(f"[!] Uyarı: Canlı model listelenemedi. Varsayılan '{self.selected_model}' kullanılacak.")
                
        except Exception as e:
            print(f"[!] Model listeleme hatası: {e}. Varsayılan modele geçiliyor...")
            self.selected_model = "models/gemini-flash-latest"

    def retrieve_all_relevant_context(self, topic, top_k=6):
        """
        Makale konusuyla ilişkili en önemli Obsidian notlarını (en fazla 6 adet) vektörel olarak çeker.
        """
        formatted_query = f"query: {topic}"
        query_vector = self.embed_model.encode(formatted_query).tolist()

        results = self.collection.query(
            query_embeddings=[query_vector],
            n_results=top_k
        )
        
        retrieved_documents = []
        sources_info = []
        
        if results and results['documents'] and len(results['documents'][0]) > 0:
            for i in range(len(results['documents'][0])):
                text = results['documents'][0][i]
                metadata = results['metadatas'][0][i]
                score = results['distances'][0][i]
                similarity = 1 - (score / 2)
                
                # Akademik derinlik için eşiği %65'e çekiyoruz ki ilişkili kavramlar da gelsin
                if similarity >= 0.65:
                    retrieved_documents.append({
                        "title": metadata['title'],
                        "content": text,
                        "links": metadata.get('links', '')
                    })
                    sources_info.append(f"[[{metadata['title']}]] (%{similarity*100:.1f} Alaka)")
                    
        return retrieved_documents, sources_info

    def write_academic_paper(self, topic, outline_focus=None):
        """
        Çekilen notları sentezleyerek üst düzey, referanslı bir makale taslağı yazar.
        """
        docs, sources_info = self.retrieve_all_relevant_context(topic, top_k=6)
        
        if not docs:
            return "Kütüphanenizde bu konuyla ilişkilendirilebilecek yeterli not bulunamadı.", []
            
        # Bağlam dizesini oluştur
        context_str = ""
        for doc in docs:
            context_str += f"\n--- REFERANS NOTU: [[{doc['title']}]] ---\n"
            context_str += f"İlişkili Kavramlar: {doc['links']}\n"
            context_str += f"Not İçeriği:\n{doc['content']}\n"
            context_str += "----------------------------------------\n"

        # Sistem Yönergesi (Akademik Yazar Ruhu)
        system_instruction = """Sen İlahiyat, İslam Felsefesi ve Dijital Beşerî Bilimler (Digital Humanities) alanında uzmanlaşmış, uluslararası hakemli dergilerde (Peer-Reviewed) yayın yapan kıdemli bir akademisyensin.
Görevin, sana sunulan 'KAYNAK NOTLAR' verilerini kullanarak, akademik metodolojiye tamamen uygun, derinlikli ve yapılandırılmış bir makale taslağı yazmaktır.

Senden beklenen akademik format:
1. Başlık: Konuya uygun, akademik ağırlığı olan şık bir başlık.
2. Özet (Abstract): Makalenin amacını, yöntemini ve ana argümanını özetleyen Türkçe ve İngilizce (yaklaşık 150 kelime) özet bölümü.
3. Giriş (Introduction): Konunun önemi, literatürdeki yeri ve problem durumu.
4. Gelişme Bölümleri (Ana Gövde): Çekilen kaynak notları tematik başlıklar altında birleştiren, derin analizler içeren alt bölümler.
5. Sonuç (Conclusion): Elde edilen felsefi ve teolojik çıkarımların sentezi.
6. Kaynakça (References): Sadece kullanılan kaynak notların isimlerinden oluşan akademik kaynakça listesi.

Uyman Gereken Altın Akademik Kurallar:
- SADECE sana sunulan kaynak notlardaki felsefi argümanları, kavramları ve verileri kullan. Dışarıdan genel kültür veya bilgi ekleme.
- Makale metninin içinde geçen her felsefi iddia, tanım veya aktarımdan sonra, hangi nottan alındığını çift köşeli parantez referansıyla göster. (Örn: "...duyusal verilerin insanı aldatabileceğini belirterek hissebileceğimiz yanılgılara dikkat çeker [[Hissiyyât]].")
- Dilin son derece ağırbaşlı, akademik, analitik ve saygın olmalıdır. Türkçe gramer kurallarına kusursuz uymalıdır.
- Kavramlar arasındaki bağlantıları (örneğin Şüphe ile Yakîn veya Zaruriyyât ile Tasavvuf arasındaki geçişleri) entelektüel bir derinlikle işle.
"""

        user_message = f"""Makale Konusu: {topic}
Ek Odak Noktası / Kullanıcı Direktifi: {outline_focus if outline_focus else 'Genel kavramsal analiz ve metodolojik değerlendirme.'}

Sana Sunulan Kaynak Notlar (Bu notların dışına çıkma):
==================================================
{context_str}
==================================================

Yukarıdaki kaynak notların tamamını akademik bir disiplinle sentezleyerek, iç atıfları (wikilink formatında) eksiksiz olan muazzam bir makale taslağı kaleme al:"""

        print(f"\n[*] {self.selected_model} modeli kullanılarak makale taslağınız kaleme alınıyor... (Bu işlem 15-20 saniye sürebilir)")
        
        model = genai.GenerativeModel(
            model_name=self.selected_model,
            system_instruction=system_instruction
        )
        
        response = model.generate_content(
            contents=user_message,
            generation_config=genai.types.GenerationConfig(temperature=0.2)
        )
        
        return response.text, sources_info


# =====================================================================
# ETKİLEŞİMLİ ÇALIŞTIRMA KONSOLU
# =====================================================================
if __name__ == "__main__":
    print("\n" + "="*70)
    print("✍️  İMAM GAZALİ KÜLLİYATI AKADEMİK CO-WRITER (YAZAR ASİSTANI)")
    print("="*70)
    
    try:
        writer = GazaliAcademicCoWriter()
        print("\n[+] Akademik asistanınız hazır!")
        
        while True:
            print("\n" + "-"*50)
            konu = input("📝 Makale konusunu / anahtar kelimeleri girin (Çıkış için 'q'): ").strip()
            
            if not konu:
                continue
            if konu.lower() in ['q', 'exit', 'quit']:
                print("\n[*] Oturum kapatılıyor. Başarılar dileriz!")
                break
                
            odak = input("🎯 Özel olarak odaklanılmasını istediğiniz bir açı var mı? (Enter ile geçebilirsiniz): ").strip()
            
            try:
                taslak, kaynaklar = writer.write_academic_paper(konu, odak)
                
                print("\n📚 [BU MAKALE İÇİN KULLANILAN OBSIDIAN NOTLARI]")
                for k in kaynaklar:
                    print(f"   • {k}")
                
                print("\n💾 Hangi formatta kaydetmek istersiniz?")
                print("   1. Markdown (.md) - Obsidian kasanıza doğrudan eklemek için")
                print("   2. Word (.docx) - Microsoft Word belgesi olarak kaydetmek için")
                secim = input("👉 Seçiminiz (1 veya 2): ").strip()
                
                # Dosya adını temizleyelim
                safe_title = "".join([c for c in konu if c.isalpha() or c.isdigit() or c==' ']).rstrip()
                safe_title = safe_title.replace(" ", "_")[:30]
                
                if secim == "2":
                    try:
                        import docx
                        doc_file = f"Gazali_Makale_{safe_title}.docx"
                        doc = docx.Document()
                        doc.add_heading(f"Gazali Araştırmaları: {konu}", level=1)
                        
                        # Markdown metnini paragraflara bölerek Word'e yazalım
                        paragraphs = taslak.split("\n")
                        for p in paragraphs:
                            p_clean = p.strip()
                            if p_clean.startswith("### "):
                                doc.add_heading(p_clean.replace("### ", ""), level=3)
                            elif p_clean.startswith("## "):
                                doc.add_heading(p_clean.replace("## ", ""), level=2)
                            elif p_clean:
                                doc.add_paragraph(p_clean)
                                
                        doc.save(doc_file)
                        print(f"\n[🎉] Tebrikler! Makaleniz başarıyla kaydedildi: {doc_file}")
                    except ImportError:
                        print("\n[!] Hata: Bilgisayarınızda 'python-docx' kütüphanesi kurulu değil!")
                        print("[i] Terminalde 'pip install python-docx' komutunu çalıştırarak kurabilirsiniz.")
                        print("[*] Makale taslağı şimdilik zorunlu olarak Markdown (.md) formatında kaydediliyor...")
                        secim = "1"
                
                if secim == "1" or secim != "2":
                    md_file = f"Gazali_Makale_{safe_title}.md"
                    with open(md_file, "w", encoding="utf-8") as f:
                        f.write(taslak)
                    print(f"\n[🎉] Tebrikler! Makaleniz başarıyla kaydedildi: {md_file}")
                    print("[i] Bu dosyayı doğrudan Obsidian kasanızın içine kopyalayarak grafik görünümünde anında izleyebilirsiniz!")
                    
            except Exception as e:
                print(f"\n[!] Makale taslağı üretilirken bir hata oluştu: {e}")
                
    except Exception as e:
        print(f"\n[!] Sistem başlatılamadı: {e}")
